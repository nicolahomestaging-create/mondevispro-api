"""
MonDevisPro API
Génère des devis et factures PDF + Word professionnels
Version 3.0.0
"""

from fastapi import FastAPI, HTTPException, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import uuid
import resend
import json
import re
from datetime import datetime, timedelta
import requests
from io import BytesIO
from openai import OpenAI  # Gardé pour Whisper uniquement
from anthropic import Anthropic  # Claude Sonnet pour le chat

# PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor, white
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

# Word
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
# Supabase Storage
from supabase import create_client, Client

app = FastAPI(
    title="MonDevisPro API",
    description="API de génération de devis et factures PDF + Word",
    version="3.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

PDF_FOLDER = "generated_pdfs"
os.makedirs(PDF_FOLDER, exist_ok=True)

# Configuration Supabase Storage
# Essayer plusieurs noms de variables possibles (Railway peut utiliser différents préfixes)
SUPABASE_URL = (
    os.getenv("SUPABASE_URL") or 
    os.getenv("RAILWAY_SUPABASE_URL") or
    os.getenv("DATABASE_URL") or  # Parfois Railway utilise DATABASE_URL
    ""
)
SUPABASE_SERVICE_KEY = (
    os.getenv("SUPABASE_SERVICE_KEY") or 
    os.getenv("RAILWAY_SUPABASE_SERVICE_KEY") or
    os.getenv("SUPABASE_SERVICE_ROLE_KEY") or
    ""
)

# Debug: Afficher TOUTES les variables d'environnement qui contiennent "SUPABASE"
print("=== DEBUG ENV VARIABLES ===")
all_env = {k: v[:20] + "..." if v and len(v) > 20 else v for k, v in os.environ.items() if "SUPABASE" in k.upper() or "DATABASE" in k.upper()}
for key, value in all_env.items():
    print(f"{key}: {value}")
print("==========================")

print(f"=== SUPABASE CONFIG ===")
print(f"SUPABASE_URL (env): {'OUI' if os.getenv('SUPABASE_URL') else 'NON'}")
print(f"SUPABASE_SERVICE_KEY (env): {'OUI' if os.getenv('SUPABASE_SERVICE_KEY') else 'NON'}")
print(f"URL finale: {SUPABASE_URL[:50] if SUPABASE_URL else 'VIDE'}...")
print(f"KEY finale: {SUPABASE_SERVICE_KEY[:20] if SUPABASE_SERVICE_KEY else 'VIDE'}...")
print(f"Longueur URL: {len(SUPABASE_URL) if SUPABASE_URL else 0}")
print(f"Longueur KEY: {len(SUPABASE_SERVICE_KEY) if SUPABASE_SERVICE_KEY else 0}")
print(f"=======================")

# Initialiser le client Supabase UNE SEULE FOIS
supabase_client: Optional[Client] = None
if SUPABASE_URL and SUPABASE_SERVICE_KEY:
    try:
        supabase_client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
        print("✅ Supabase client créé")
        
        # Vérifier que le bucket 'documents' existe
        try:
            buckets = supabase_client.storage.list_buckets()
            bucket_names = [b.name for b in buckets]
            if 'documents' not in bucket_names:
                print("⚠️ ATTENTION: Le bucket 'documents' n'existe pas dans Supabase Storage!")
                print(f"   Buckets disponibles: {bucket_names}")
            else:
                print("✅ Bucket 'documents' trouvé")
        except Exception as e:
            print(f"⚠️ Erreur lors de la vérification des buckets: {e}")
    except Exception as e:
        print(f"❌ Erreur lors de la création du client Supabase: {e}")
        supabase_client = None
else:
    print("❌ Supabase non configuré - variables d'environnement manquantes")

# =============================================================================
# CONFIGURATION ANTHROPIC (Claude Sonnet)
# =============================================================================
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

anthropic_client = None
if ANTHROPIC_API_KEY:
    try:
        anthropic_client = Anthropic(api_key=ANTHROPIC_API_KEY)
        print("✅ Anthropic client (Claude Sonnet) configuré")
    except Exception as e:
        print(f"❌ Erreur configuration Anthropic: {e}")
else:
    print("⚠️ ANTHROPIC_API_KEY non configurée - Claude désactivé")

# Garder OpenAI pour Whisper uniquement
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
openai_whisper_client = None
if OPENAI_API_KEY:
    try:
        openai_whisper_client = OpenAI(api_key=OPENAI_API_KEY)
        print("✅ OpenAI client (Whisper) configuré")
    except Exception as e:
        print(f"❌ Erreur configuration OpenAI: {e}")

def upload_to_supabase(filepath: str, filename: str) -> str:
    """Upload un fichier sur Supabase Storage et retourne l'URL publique"""
    if not supabase_client:
        print(f"⚠️ Supabase non configuré, fichier local conservé: {filename}")
        return f"/download/{filename}"
    
    try:
        # Vérifier que le fichier existe
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Le fichier {filepath} n'existe pas")
        
        file_size = os.path.getsize(filepath)
        print(f"📁 Taille du fichier {filename}: {file_size} bytes")
        
        # Lire le fichier
        with open(filepath, 'rb') as f:
            file_data = f.read()
        
        if len(file_data) == 0:
            raise ValueError(f"Le fichier {filename} est vide")
        
        print(f"📤 Début upload de {filename} ({len(file_data)} bytes)")
        
        # Déterminer le content-type
        content_type = "application/pdf" if filename.endswith('.pdf') else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        # Essayer de supprimer le fichier existant d'abord
        try:
            result = supabase_client.storage.from_('documents').remove([filename])
            print(f"🗑️  Tentative de suppression du fichier existant: {result}")
        except Exception as e:
            print(f"ℹ️  Fichier n'existe pas encore (normal): {e}")
        
        # Upload sur Supabase Storage
        # La bibliothèque supabase-py attend file_data directement, pas file_options avec upsert
        upload_response = supabase_client.storage.from_('documents').upload(
            path=filename,
            file=file_data,
            file_options={"content-type": content_type}
        )
        
        print(f"📥 Réponse upload: {upload_response}")
        print(f"📥 Type de réponse: {type(upload_response)}")
        
        # Vérifier que l'upload a réussi
        # La réponse peut être un dict avec 'error' ou une liste
        if isinstance(upload_response, dict) and upload_response.get('error'):
            error_msg = upload_response.get('error', 'Erreur inconnue')
            raise Exception(f"Erreur upload Supabase: {error_msg}")
        
        print(f"✅ Upload réussi pour {filename}")
        
        # Générer l'URL publique
        # get_public_url retourne directement une chaîne d'URL
        public_url = supabase_client.storage.from_('documents').get_public_url(filename)
        
        print(f"🔗 Type URL publique: {type(public_url)}")
        print(f"🔗 URL publique brute: {public_url}")
        
        # Convertir en string si nécessaire
        if isinstance(public_url, dict):
            public_url = public_url.get('publicUrl', '') or public_url.get('public_url', '')
        elif not isinstance(public_url, str):
            public_url = str(public_url)
        
        if not public_url or public_url == '' or public_url == 'None':
            raise Exception(f"URL publique vide ou invalide: {public_url}")
        
        print(f"✅ URL publique finale: {public_url}")
        
        # Supprimer le fichier local seulement après confirmation de l'upload
        if os.path.exists(filepath):
            try:
                os.remove(filepath)
                print(f"🗑️  Fichier local supprimé: {filepath}")
            except Exception as e:
                print(f"⚠️  Impossible de supprimer le fichier local: {e}")
        
        return public_url
        
    except FileNotFoundError as e:
        print(f"❌ Erreur fichier non trouvé: {e}")
        return f"/download/{filename}"
    except Exception as e:
        print(f"❌ Erreur upload Supabase pour {filename}: {e}")
        print(f"   Type d'erreur: {type(e).__name__}")
        import traceback
        traceback.print_exc()
        # Ne pas supprimer le fichier local en cas d'erreur
        return f"/download/{filename}"

# ==================== FONCTIONS DASHBOARD SUPABASE ====================

def get_entreprise_by_whatsapp(phone: str) -> Optional[Dict]:
    """
    Trouve l'entreprise liée à un numéro WhatsApp.
    Le numéro peut être au format:
    - whatsapp:+33605108023 (format Twilio)
    - +33605108023
    - 33605108023
    """
    if not supabase_client or not phone:
        print(f"⚠️ get_entreprise_by_whatsapp: supabase_client={bool(supabase_client)}, phone={phone}")
        return None
    
    try:
        # Normaliser le numéro (enlever whatsapp:, + et espaces)
        phone_normalized = phone.replace('whatsapp:', '').replace('+', '').strip()
        print(f"📱 Recherche entreprise pour WhatsApp: {phone} -> normalisé: {phone_normalized}")
        
        # Chercher l'entreprise par le champ whatsapp
        result = supabase_client.table('entreprises').select('*').eq('whatsapp', phone_normalized).execute()
        
        if result.data and len(result.data) > 0:
            print(f"✅ Entreprise trouvée pour WhatsApp {phone_normalized}: {result.data[0].get('nom')}")
            return result.data[0]
        
        # Essayer aussi avec le champ tel (si whatsapp non configuré)
        result = supabase_client.table('entreprises').select('*').eq('tel', phone_normalized).execute()
        
        if result.data and len(result.data) > 0:
            print(f"✅ Entreprise trouvée par tel {phone_normalized}: {result.data[0].get('nom')}")
            return result.data[0]
        
        print(f"⚠️ Aucune entreprise trouvée pour le numéro {phone_normalized}")
        return None
        
    except Exception as e:
        print(f"❌ Erreur recherche entreprise par WhatsApp: {e}")
        return None


def save_devis_to_dashboard(
    entreprise_id: str,
    numero_devis: str,
    client_nom: str,
    client_email: Optional[str],
    client_telephone: Optional[str],
    titre_projet: Optional[str],
    prestations: List[Dict],
    total_ht: float,
    total_ttc: float,
    pdf_url: Optional[str],
    word_url: Optional[str],
    remise_type: Optional[str] = None,
    remise_value: Optional[float] = None,
    delai: Optional[str] = None
) -> Optional[Dict]:
    """
    Sauvegarde un devis dans la table dashboard (même table que le site web).
    Retourne le devis créé ou None en cas d'erreur.
    """
    if not supabase_client or not entreprise_id:
        print("⚠️ Supabase non configuré ou entreprise_id manquant, devis non sauvegardé en base")
        return None
    
    try:
        # Préparer les prestations au format JSON string (comme le dashboard)
        prestations_json = json.dumps(prestations, ensure_ascii=False)
        
        devis_data = {
            'entreprise_id': entreprise_id,
            'numero_devis': numero_devis,
            'client_nom': client_nom,
            'client_email': client_email,
            'telephone_client': client_telephone,
            'titre_projet': titre_projet,
            'prestations': prestations_json,
            'total_ht': total_ht,
            'total_ttc': total_ttc,
            'statut': 'en_attente',
            'date': datetime.now().strftime('%Y-%m-%d'),
            'pdf_url': pdf_url,
            'word_url': word_url,
            'remise_type': remise_type,
            'remise_value': remise_value if remise_type and remise_value else None,
        }
        
        result = supabase_client.table('devis').insert(devis_data).execute()
        
        if result.data and len(result.data) > 0:
            print(f"✅ Devis {numero_devis} sauvegardé dans dashboard (id: {result.data[0].get('id')})")
            return result.data[0]
        else:
            print(f"⚠️ Devis inséré mais pas de données retournées")
            return None
            
    except Exception as e:
        print(f"❌ Erreur sauvegarde devis dashboard: {e}")
        import traceback
        traceback.print_exc()
        return None


def save_facture_to_dashboard(
    entreprise_id: str,
    devis_id: Optional[str],
    numero_facture: str,
    client_nom: str,
    client_email: Optional[str],
    client_telephone: Optional[str],
    client_adresse: Optional[str],
    titre_projet: Optional[str],
    prestations: List[Dict],
    total_ht: float,
    total_ttc: float,
    pdf_url: Optional[str],
    word_url: Optional[str],
    type_facture: str = 'complete',  # 'acompte' ou 'complete'
    remise_type: Optional[str] = None,
    remise_value: Optional[float] = None,
    tva_taux: float = 20.0,
    solde_a_payer: Optional[float] = None
) -> Optional[Dict]:
    """
    Sauvegarde une facture dans la table dashboard (même table que le site web).
    Retourne la facture créée ou None en cas d'erreur.
    """
    if not supabase_client or not entreprise_id:
        print("⚠️ Supabase non configuré ou entreprise_id manquant, facture non sauvegardée en base")
        return None
    
    try:
        # Préparer les prestations au format JSON string
        prestations_json = json.dumps(prestations, ensure_ascii=False)
        
        facture_data = {
            'entreprise_id': entreprise_id,
            'numero_facture': numero_facture,
            'client_nom': client_nom,
            'client_email': client_email,
            'client_telephone': client_telephone,
            'client_adresse': client_adresse,
            'titre_projet': titre_projet,
            'prestations': prestations_json,
            'total_ht': total_ht,
            'total_ttc': total_ttc,
            'statut': 'en_attente',
            'date': datetime.now().strftime('%Y-%m-%d'),
            'pdf_url': pdf_url,
            'word_url': word_url,
            'type_facture': type_facture,
            'remise_type': remise_type,
            'remise_value': remise_value if remise_type and remise_value else None,
            'tva_taux': tva_taux,
            'solde_a_payer': solde_a_payer,
        }
        
        # Ajouter devis_id si fourni
        if devis_id:
            facture_data['devis_id'] = devis_id
        
        result = supabase_client.table('factures').insert(facture_data).execute()
        
        if result.data and len(result.data) > 0:
            print(f"✅ Facture {numero_facture} sauvegardée dans dashboard (id: {result.data[0].get('id')})")
            return result.data[0]
        else:
            print(f"⚠️ Facture insérée mais pas de données retournées")
            return None
            
    except Exception as e:
        print(f"❌ Erreur sauvegarde facture dashboard: {e}")
        import traceback
        traceback.print_exc()
        return None


def get_devis_by_numero(numero_devis: str, entreprise_id: Optional[str] = None) -> Optional[Dict]:
    """
    Récupère un devis par son numéro.
    """
    if not supabase_client or not numero_devis:
        return None
    
    try:
        query = supabase_client.table('devis').select('*').eq('numero_devis', numero_devis)
        
        if entreprise_id:
            query = query.eq('entreprise_id', entreprise_id)
        
        result = query.execute()
        
        if result.data and len(result.data) > 0:
            print(f"✅ Devis {numero_devis} trouvé")
            return result.data[0]
        
        print(f"⚠️ Devis {numero_devis} non trouvé")
        return None
        
    except Exception as e:
        print(f"❌ Erreur recherche devis: {e}")
        return None


# Couleurs par défaut (utilisées si couleur_pdf n'est pas défini)
COULEUR_DEFAUT = '#2F665B'
BLEU_CLAIR = HexColor('#3498db')
GRIS_FONCE = HexColor('#2c3e50')
GRIS_CLAIR = HexColor('#ecf0f1')
GRIS_TEXTE = HexColor('#555555')


# ==================== MODÈLES ====================

class Prestation(BaseModel):
    description: str
    quantite: float
    unite: str
    prix_unitaire: float
    tva_taux: Optional[float] = None  # Taux TVA par prestation
    description_detaillee: Optional[str] = None  # Description longue (sous la description principale)
    notes: Optional[str] = None  # Notes en italique

class LigneFinale(BaseModel):
    """Ligne finale du devis figé avec montant HT après remise"""
    description: str
    quantite: float = 1
    unite: str = "u"
    ht_apres_remise: float  # Montant HT après remise pour cette ligne
    tva_taux: float = 20.0  # Taux TVA pour cette ligne
    description_detaillee: Optional[str] = None  # Description longue
    notes: Optional[str] = None  # Notes en italique

class Entreprise(BaseModel):
    nom: str
    gerant: Optional[str] = ""
    siret: str
    adresse: str
    cp_ville: str
    tel: str
    email: str = ""
    logo_url: Optional[str] = None
    tva_taux: Optional[float] = 20.0
    mention_legale_tva: Optional[str] = ""
    conditions_paiement: Optional[str] = "30% à la commande, solde à réception"
    delai_validite: Optional[int] = 30
    forme_juridique: Optional[str] = None  # Ne pas forcer auto-entrepreneur par défaut
    capital_social: Optional[str] = ""
    rcs: Optional[str] = ""
    tva_intracommunautaire: Optional[str] = ""
    couleur_pdf: Optional[str] = None

class Client(BaseModel):
    nom: str
    adresse: Optional[str] = ""
    cp_ville: Optional[str] = ""
    tel: Optional[str] = ""
    email: Optional[str] = ""

class DevisRequest(BaseModel):
    entreprise: Entreprise
    client: Client
    prestations: List[Prestation]
    tva_taux: float = 20.0
    conditions_paiement: str = "30% a la commande, solde a reception"
    delai_realisation: str = "A definir"
    validite_jours: int = 30
    remise_type: Optional[str] = None  # "pourcentage" ou "fixe"
    remise_valeur: Optional[float] = 0
    acompte_pourcentage: Optional[float] = 0
    numero_devis: Optional[str] = None  # Numero de devis fourni par le client (OBLIGATOIRE)

class DevisDataFromAI(BaseModel):
    client_nom: str
    client_adresse: Optional[str] = ""
    client_email: Optional[str] = ""
    client_telephone: Optional[str] = ""
    titre_projet: Optional[str] = ""
    prestations: Optional[List[Prestation]] = None
    prestations_json: Optional[str] = None  # Alternative: prestations comme string JSON
    delai: Optional[str] = "A definir"
    remise_type: Optional[str] = None
    remise_valeur: Optional[float] = 0
    acompte_pourcentage: Optional[float] = 0

class DevisRequestSimple(BaseModel):
    entreprise: Entreprise
    devis_data: DevisDataFromAI
    validite_jours: int = 30
    phone: Optional[str] = None  # Numéro WhatsApp pour sauvegarde automatique dashboard

class RIB(BaseModel):
    iban: Optional[str] = ""
    bic: Optional[str] = ""
    titulaire: Optional[str] = ""
    
class FactureRequest(BaseModel):
    entreprise: Entreprise
    client: Client
    prestations: Optional[List[Prestation]] = None  # Optionnel si prestations_json fourni
    tva_taux: float = 20.0
    numero_devis_origine: Optional[str] = None
    numero_facture: Optional[str] = None  # Numéro de facture fourni par le frontend
    date_echeance_jours: int = 30
    mention_legale_tva: Optional[str] = ""
    rib: Optional[RIB] = None
    remise_type: Optional[str] = None  # "pourcentage" ou "montant"
    remise_valeur: Optional[float] = 0
    statut: Optional[str] = "en_attente"  # "en_attente", "payee", etc.
    total_ht: Optional[float] = None  # Total HT pour factures d'acompte
    total_ttc: Optional[float] = None  # Total TTC pour factures d'acompte
    total_ht_devis: Optional[float] = None  # Total HT du devis (avec remise)
    total_ttc_devis: Optional[float] = None  # Total TTC du devis (avec remise)
    prestations_json: Optional[str] = None  # Prestations comme string JSON encodée URL
    is_facture_acompte: Optional[bool] = False  # Flag pour factures d'acompte
    taux_acompte: Optional[float] = None  # Taux d'acompte en pourcentage
    acompte_ttc_deja_facture: Optional[float] = None  # Montant TTC des acomptes déjà versés
    acompte_references: Optional[List[str]] = None  # Numéros des factures d'acompte
    lignes_finales_devis: Optional[List[LigneFinale]] = None  # Lignes finales du devis figé (priorité sur prestations)
    phone: Optional[str] = None  # Numéro WhatsApp pour sauvegarde automatique dashboard


# ==================== FONCTIONS UTILITAIRES ====================

def get_couleur_principale(data) -> HexColor:
    """Récupère la couleur principale depuis couleur_pdf ou utilise la couleur par défaut"""
    couleur_hex = data.entreprise.couleur_pdf if data.entreprise.couleur_pdf else COULEUR_DEFAUT
    # S'assurer que la couleur commence par #
    if not couleur_hex.startswith('#'):
        couleur_hex = '#' + couleur_hex
    try:
        return HexColor(couleur_hex)
    except:
        # En cas d'erreur, utiliser la couleur par défaut
        return HexColor(COULEUR_DEFAUT)

def hex_to_rgb(hex_color: str) -> tuple:
    """Convertit une couleur hex (#RRGGBB) en tuple RGB (r, g, b)"""
    # Enlever le # si présent
    hex_color = hex_color.lstrip('#')
    try:
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    except:
        # En cas d'erreur, retourner la couleur par défaut
        hex_default = COULEUR_DEFAUT.lstrip('#')
        return tuple(int(hex_default[i:i+2], 16) for i in (0, 2, 4))

def get_couleur_principale_rgb(data) -> RGBColor:
    """Récupère la couleur principale au format RGBColor pour Word"""
    couleur_hex = data.entreprise.couleur_pdf if data.entreprise.couleur_pdf else COULEUR_DEFAUT
    # S'assurer que la couleur commence par #
    if not couleur_hex.startswith('#'):
        couleur_hex = '#' + couleur_hex
    r, g, b = hex_to_rgb(couleur_hex)
    return RGBColor(r, g, b)

def get_couleur_principale_hex_string(data) -> str:
    """Récupère la couleur principale au format hex string (sans #) pour Word set_cell_shading"""
    couleur_hex = data.entreprise.couleur_pdf if data.entreprise.couleur_pdf else COULEUR_DEFAUT
    # Enlever le # si présent
    return couleur_hex.lstrip('#')

def telecharger_logo(logo_url: str) -> Optional[ImageReader]:
    try:
        if not logo_url or logo_url.strip() == "":
            return None
        response = requests.get(logo_url, timeout=10)
        if response.status_code == 200:
            image_data = BytesIO(response.content)
            return ImageReader(image_data)
    except Exception as e:
        print(f"Erreur téléchargement logo: {e}")
    return None

def telecharger_logo_bytes(logo_url: str) -> Optional[BytesIO]:
    """Télécharge le logo et retourne les bytes pour Word"""
    try:
        if not logo_url or logo_url.strip() == "":
            return None
        response = requests.get(logo_url, timeout=10)
        if response.status_code == 200:
            return BytesIO(response.content)
    except Exception as e:
        print(f"Erreur téléchargement logo: {e}")
    return None

def tronquer_texte(texte: str, max_chars: int) -> str:
    if not texte:
        return ""
    if len(texte) <= max_chars:
        return texte
    return texte[:max_chars-3] + "..."

def decouper_texte_en_lignes(texte: str, max_chars: int = 45) -> list:
    """Découpe un texte long en plusieurs lignes sans couper les mots"""
    if not texte:
        return []
    
    lignes = []
    mots = texte.split()
    ligne_courante = ""
    
    for mot in mots:
        test_ligne = (ligne_courante + " " + mot).strip() if ligne_courante else mot
        if len(test_ligne) <= max_chars:
            ligne_courante = test_ligne
        else:
            if ligne_courante:
                lignes.append(ligne_courante)
            # Si le mot seul est trop long, on le tronque
            if len(mot) > max_chars:
                lignes.append(mot[:max_chars-3] + "...")
                ligne_courante = ""
            else:
                ligne_courante = mot
    
    if ligne_courante:
        lignes.append(ligne_courante)
    
    return lignes

def formater_adresse_complete(adresse: str, cp_ville: str) -> str:
    parties = []
    if adresse and adresse.strip():
        parties.append(adresse.strip())
    if cp_ville and cp_ville.strip():
        parties.append(cp_ville.strip())
    return ", ".join(parties) if parties else ""


# ==================== GÉNÉRATION PDF ====================

def dessiner_bloc_emetteur(c, width, height, data, y_position):
    c.setFillColor(GRIS_CLAIR)
    c.roundRect(15*mm, y_position - 32*mm, 85*mm, 38*mm, 3*mm, fill=True, stroke=False)
    
    c.setFillColor(get_couleur_principale(data))
    c.setFont("Helvetica-Bold", 10)
    c.drawString(20*mm, y_position, "ÉMETTEUR")
    
    c.setFillColor(GRIS_FONCE)
    c.setFont("Helvetica", 9)
    y_text = y_position - 5*mm
    
    c.drawString(20*mm, y_text, tronquer_texte(data.entreprise.nom, 40))
    
    adresse = data.entreprise.adresse if data.entreprise.adresse else ""
    cp_ville = data.entreprise.cp_ville if data.entreprise.cp_ville else ""
    
    ligne_y = y_text - 5*mm
    
    if len(adresse) <= 35:
        if adresse:
            c.drawString(20*mm, ligne_y, adresse)
            ligne_y -= 5*mm
    else:
        mots = adresse.split()
        ligne1 = ""
        ligne2 = ""
        for mot in mots:
            if len(ligne1 + " " + mot) <= 35:
                ligne1 = (ligne1 + " " + mot).strip()
            else:
                ligne2 = (ligne2 + " " + mot).strip()
        c.drawString(20*mm, ligne_y, ligne1)
        ligne_y -= 5*mm
        if ligne2:
            c.drawString(20*mm, ligne_y, ligne2)
            ligne_y -= 5*mm
    
    if cp_ville:
        c.drawString(20*mm, ligne_y, cp_ville)
        ligne_y -= 5*mm
    
    c.drawString(20*mm, ligne_y, f"Tél : {data.entreprise.tel}")
    ligne_y -= 5*mm
    c.drawString(20*mm, ligne_y, f"Email : {tronquer_texte(data.entreprise.email, 35)}")
    ligne_y -= 5*mm
    c.drawString(20*mm, ligne_y, f"SIRET : {data.entreprise.siret}")


def dessiner_bloc_client(c, width, height, data, y_position):
    c.setFillColor(GRIS_CLAIR)
    c.roundRect(110*mm, y_position - 32*mm, 85*mm, 38*mm, 3*mm, fill=True, stroke=False)
    
    c.setFillColor(get_couleur_principale(data))
    c.setFont("Helvetica-Bold", 10)
    c.drawString(115*mm, y_position, "DESTINATAIRE")
    
    c.setFillColor(GRIS_FONCE)
    c.setFont("Helvetica", 9)
    y_text = y_position - 5*mm
    
    c.drawString(115*mm, y_text, data.client.nom)
    ligne_y = y_text - 5*mm
    
    if data.client.adresse:
        adresse = data.client.adresse
        if len(adresse) <= 35:
            c.drawString(115*mm, ligne_y, adresse)
            ligne_y -= 5*mm
        else:
            mots = adresse.split()
            ligne1 = ""
            ligne2 = ""
            for mot in mots:
                if len(ligne1 + " " + mot) <= 35:
                    ligne1 = (ligne1 + " " + mot).strip()
                else:
                    ligne2 = (ligne2 + " " + mot).strip()
            c.drawString(115*mm, ligne_y, ligne1)
            ligne_y -= 5*mm
            if ligne2:
                c.drawString(115*mm, ligne_y, ligne2)
                ligne_y -= 5*mm
    
    if data.client.cp_ville:
        c.drawString(115*mm, ligne_y, data.client.cp_ville)
        ligne_y -= 5*mm
    
    if data.client.tel:
        c.drawString(115*mm, ligne_y, f"Tél : {data.client.tel}")
        ligne_y -= 5*mm
    
    if data.client.email:
        c.drawString(115*mm, ligne_y, f"Email : {data.client.email}")


def dessiner_en_tete_page(c, width, height, data, numero_devis, logo, date_validite):
    """Dessine l'en-tête de page (pour la première page et les pages suivantes)"""
    print(f"🔍 dessiner_en_tete_page - numero_devis reçu: '{numero_devis}'")
    c.setFillColor(get_couleur_principale(data))
    c.rect(0, height - 45*mm, width, 45*mm, fill=True, stroke=False)
    
    text_start_x = 15*mm
    
    if logo:
        try:
            logo_size = 30*mm
            c.drawImage(logo, 15*mm, height - 40*mm, width=logo_size, height=logo_size, preserveAspectRatio=True, mask='auto')
            text_start_x = 50*mm
        except Exception as e:
            print(f"Erreur logo: {e}")
    
    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 18)
    c.drawString(text_start_x, height - 18*mm, tronquer_texte(data.entreprise.nom.upper(), 30))
    
    if data.entreprise.gerant and data.entreprise.gerant.strip():
        c.setFont("Helvetica", 9)
        c.drawString(text_start_x, height - 26*mm, f"Gérant : {data.entreprise.gerant}")
    
    c.setFont("Helvetica-Bold", 28)
    c.drawRightString(width - 20*mm, height - 18*mm, "DEVIS")
    c.setFont("Helvetica", 11)
    c.drawRightString(width - 20*mm, height - 28*mm, f"N° {numero_devis}")
    c.setFont("Helvetica", 9)
    c.drawRightString(width - 20*mm, height - 36*mm, f"Date : {datetime.now().strftime('%d/%m/%Y')}")


def dessiner_totaux(c, width, y_totaux, total_ht, total_ht_avant_acompte, total_acompte, remise, tva_taux, total_ht_final, total_ttc, data):
    """
    Dessine les totaux pour un devis avec affichage de la remise si présente
    """
    x_label = 130*mm
    x_value = width - 18*mm
    c.setFillColor(GRIS_FONCE)
    c.setFont("Helvetica", 10)
    
    y_offset = 0
    
    # Récupérer les informations de remise depuis data
    remise_type = getattr(data, 'remise_type', None)
    remise_valeur_raw = getattr(data, 'remise_valeur', None)
    
    # Convertir remise_valeur en nombre
    remise_valeur = 0
    if remise_valeur_raw is not None:
        try:
            remise_valeur = float(remise_valeur_raw)
        except (ValueError, TypeError):
            remise_valeur = 0
    
    # Normaliser remise_type
    if remise_type:
        remise_type = str(remise_type).strip()
        if remise_type == "" or remise_type.lower() == "none":
            remise_type = None
    
    # Calculer la remise totale à partir de data si nécessaire
    remise_totale = remise
    if remise_totale == 0 and remise_type and remise_valeur > 0:
        if remise_type == "pourcentage":
            remise_totale = total_ht_avant_acompte * (remise_valeur / 100)
        elif remise_type in ["montant", "fixe"]:
            remise_totale = remise_valeur
    
    # Total HT (avant remise si remise présente)
    if remise_totale > 0:
        c.drawString(x_label, y_totaux, "Total HT avant remise")
    else:
        c.drawString(x_label, y_totaux, "Total HT")
    c.drawRightString(x_value, y_totaux, f"{total_ht_avant_acompte:.2f} €")
    y_offset = 6*mm
    
    # Afficher la remise si elle existe
    if remise_totale > 0:
        if remise_type == "pourcentage" and remise_valeur > 0:
            c.drawString(x_label, y_totaux - y_offset, f"Remise ({remise_valeur}%)")
        else:
            c.drawString(x_label, y_totaux - y_offset, "Remise")
        
        c.setFillColor(HexColor('#e74c3c'))
        c.drawRightString(x_value, y_totaux - y_offset, f"-{remise_totale:.2f} €")
        c.setFillColor(GRIS_FONCE)
        y_offset += 6*mm
        
        # Total HT après remise
        c.drawString(x_label, y_totaux - y_offset, "Total HT après remise")
        c.drawRightString(x_value, y_totaux - y_offset, f"{total_ht_final:.2f} €")
        y_offset += 6*mm
    
    # Afficher l'acompte si présent
    if total_acompte > 0:
        c.drawString(x_label, y_totaux - y_offset, "Acompte déduit")
        c.setFillColor(HexColor('#e74c3c'))
        c.drawRightString(x_value, y_totaux - y_offset, f"-{total_acompte:.2f} €")
        c.setFillColor(GRIS_FONCE)
        y_offset += 6*mm
    
    # Calculer TVA par taux à partir des prestations
    tva_par_taux = {}
    for prestation in data.prestations:
        total_ligne = prestation.quantite * prestation.prix_unitaire
        if total_ligne > 0:  # Ignorer les acomptes
            # Utiliser le taux de TVA de la prestation si disponible, sinon le taux global
            # IMPORTANT : 0 est une valeur valide (ne pas utiliser "or" qui remplacerait 0)
            presta_tva = getattr(prestation, 'tva_taux', None)
            taux = presta_tva if presta_tva is not None else tva_taux
            if taux not in tva_par_taux:
                tva_par_taux[taux] = 0
            # Appliquer la remise proportionnellement si nécessaire
            if remise_totale > 0:
                ratio_remise = (total_ht_avant_acompte - remise_totale) / total_ht_avant_acompte if total_ht_avant_acompte > 0 else 1
                total_ligne_apres_remise = total_ligne * ratio_remise
            else:
                total_ligne_apres_remise = total_ligne
            # Déduire l'acompte si présent
            total_ligne_final = total_ligne_apres_remise
            tva_par_taux[taux] += total_ligne_final * (taux / 100)
    
    # Si pas de prestations avec TVA, utiliser le calcul simple
    if not tva_par_taux:
        montant_tva_total = total_ht_final * (tva_taux / 100)
        if tva_taux > 0:
            tva_par_taux[tva_taux] = montant_tva_total
        else:
            tva_par_taux[0] = 0
    
    # Recalculer le total TTC à partir de la TVA par taux calculée
    montant_tva_total_calcule = sum(tva_par_taux.values())
    total_ttc_recalcule = total_ht_final + montant_tva_total_calcule
    
    # Afficher TVA par taux
    taux_affiches = False
    for taux in sorted(tva_par_taux.keys(), reverse=True):
        montant = tva_par_taux[taux]
        if taux > 0:
            c.drawString(x_label, y_totaux - y_offset, f"TVA ({taux}%)")
            c.drawRightString(x_value, y_totaux - y_offset, f"{montant:.2f} €")
            y_offset += 6*mm
            taux_affiches = True
    
    # Afficher "TVA non applicable" seulement si aucun taux > 0
    if not taux_affiches:
            c.setFont("Helvetica-Oblique", 8)
            c.drawString(x_label, y_totaux - y_offset, "TVA non applicable")
            c.setFont("Helvetica", 10)
            y_offset += 6*mm
    
    # Total TTC avec encadré coloré (utiliser le total_ttc recalculé)
    c.setFillColor(get_couleur_principale(data))
    c.roundRect(x_label - 5*mm, y_totaux - y_offset - 8*mm, 68*mm, 10*mm, 2*mm, fill=True, stroke=False)
    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 11)
    c.drawString(x_label, y_totaux - y_offset - 5*mm, "TOTAL TTC")
    c.drawRightString(x_value, y_totaux - y_offset - 5*mm, f"{total_ttc_recalcule:.2f} €")
    
    return y_totaux - y_offset - 8*mm  # Retourner la position Y finale


def dessiner_lignes_prestations(c, width, prestations, y_table, data, index_debut=0):
    """Dessine les lignes de prestations (en-tête + lignes) et retourne la position Y finale et les totaux calculés"""
    # En-tête du tableau
    c.setFillColor(get_couleur_principale(data))
    c.rect(15*mm, y_table, width - 30*mm, 10*mm, fill=True, stroke=False)
    
    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(18*mm, y_table + 3*mm, "Description")
    c.drawString(95*mm, y_table + 3*mm, "Qté")
    c.drawString(108*mm, y_table + 3*mm, "Unité")
    c.drawString(125*mm, y_table + 3*mm, "P.U. HT")
    c.drawString(150*mm, y_table + 3*mm, "TVA")
    c.drawRightString(width - 18*mm, y_table + 3*mm, "Total HT")
    
    y_ligne = y_table - 2*mm
    total_ht_avant_acompte = 0
    total_acompte = 0
    
    # Largeur max pour les descriptions (ne pas dépasser colonne Qté à 95mm)
    MAX_DESC_CHARS = 42  # ~75mm de large avec police 9
    MAX_DETAIL_CHARS = 40  # Pour les sous-lignes en police 7
    
    # Dessiner les lignes
    for i, prestation in enumerate(prestations):
        total_ligne = prestation.quantite * prestation.prix_unitaire
        
        # Séparer les prestations positives et les acomptes (négatifs)
        if total_ligne >= 0:
            total_ht_avant_acompte += total_ligne
        else:
            total_acompte += abs(total_ligne)
        
        # Récupérer les textes
        description_principale = getattr(prestation, 'description', '') or ''
        description_detaillee = getattr(prestation, 'description_detaillee', '') or ''
        notes = getattr(prestation, 'notes', '') or ''
        
        # Découper les textes en lignes
        lignes_desc_principale = decouper_texte_en_lignes(description_principale, MAX_DESC_CHARS)
        lignes_desc_detaillee = decouper_texte_en_lignes(description_detaillee, MAX_DETAIL_CHARS)
        lignes_notes = decouper_texte_en_lignes(notes, MAX_DETAIL_CHARS - 6)  # -6 pour "Note: "
        
        # Calculer la hauteur de ligne nécessaire
        nb_lignes_total = max(1, len(lignes_desc_principale))
        nb_lignes_total += len(lignes_desc_detaillee)
        nb_lignes_total += len(lignes_notes)
        
        # Hauteur de base + lignes supplémentaires
        if nb_lignes_total <= 1:
            hauteur_ligne = 10*mm
        else:
            hauteur_ligne = 8*mm + (nb_lignes_total * 3.5*mm)
        
        y_ligne -= hauteur_ligne
        
        # Fond alterné
        if (index_debut + i) % 2 == 0:
            c.setFillColor(HexColor('#f8f9fa'))
            c.rect(15*mm, y_ligne, width - 30*mm, hauteur_ligne, fill=True, stroke=False)
        
        # Position Y pour le texte (en haut de la cellule)
        y_text = y_ligne + hauteur_ligne - 5*mm
        
        # Description principale (peut être sur plusieurs lignes)
        c.setFillColor(GRIS_FONCE)
        c.setFont("Helvetica-Bold", 9)
        for j, ligne_desc in enumerate(lignes_desc_principale[:3]):  # Max 3 lignes
            c.drawString(18*mm, y_text, ligne_desc)
            y_text -= 3.5*mm
            if j == 0:
                c.setFont("Helvetica", 9)  # Normal après la première ligne
        
        # Description détaillée (en gris, plus petit)
        if lignes_desc_detaillee:
            c.setFont("Helvetica", 7)
            c.setFillColor(HexColor('#555555'))
            for ligne_detail in lignes_desc_detaillee[:4]:  # Max 4 lignes
                c.drawString(18*mm, y_text, ligne_detail)
                y_text -= 3*mm
        
        # Notes (en italique gris)
        if lignes_notes:
            c.setFont("Helvetica-Oblique", 7)
            c.setFillColor(HexColor('#777777'))
            for k, ligne_note in enumerate(lignes_notes[:2]):  # Max 2 lignes
                prefix = "Note: " if k == 0 else "      "
                c.drawString(18*mm, y_text, prefix + ligne_note)
                y_text -= 3*mm
        
        # Colonnes standard (alignées en haut de la cellule)
        y_colonnes = y_ligne + hauteur_ligne - 5*mm
        c.setFont("Helvetica", 9)
        c.setFillColor(GRIS_FONCE)
        c.drawString(97*mm, y_colonnes, str(prestation.quantite))
        c.drawString(108*mm, y_colonnes, getattr(prestation, 'unite', 'u') or 'u')
        c.drawString(125*mm, y_colonnes, f"{prestation.prix_unitaire:.2f} €")
        # IMPORTANT : 0 est une valeur valide pour TVA (ne pas utiliser "or")
        presta_tva_val = getattr(prestation, 'tva_taux', None)
        tva_prestation = presta_tva_val if presta_tva_val is not None else data.tva_taux
        c.drawString(150*mm, y_colonnes, f"{tva_prestation}%")
        c.drawRightString(width - 18*mm, y_colonnes, f"{total_ligne:.2f} €")
    
    y_ligne -= 5*mm
    
    # Ligne de séparation
    c.setStrokeColor(GRIS_CLAIR)
    c.setLineWidth(1)
    c.line(15*mm, y_ligne, width - 15*mm, y_ligne)
    
    return y_ligne - 10*mm, total_ht_avant_acompte, total_acompte


def dessiner_facture_depuis_lignes_finales(c, width, data, y_table, tva_taux, lignes_finales, acompte_ttc, acompte_refs):
    """
    Dessine une facture finale à partir des lignes finales du devis figé.
    Les lignes finales contiennent déjà les montants HT après remise et les TVA par ligne.
    """
    print(f"🔒 FACTURE DEPUIS LIGNES FINALES - {len(lignes_finales)} lignes")
    
    # En-tête du tableau
    c.setFillColor(get_couleur_principale(data))
    c.rect(15*mm, y_table, width - 30*mm, 10*mm, fill=True, stroke=False)
    
    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(18*mm, y_table + 3*mm, "Description")
    c.drawString(95*mm, y_table + 3*mm, "Qté")
    c.drawString(108*mm, y_table + 3*mm, "Unité")
    c.drawString(125*mm, y_table + 3*mm, "P.U. HT")
    c.drawString(150*mm, y_table + 3*mm, "TVA")
    c.drawRightString(width - 18*mm, y_table + 3*mm, "Total HT")
    
    y_ligne = y_table - 2*mm
    
    # Largeur max pour les descriptions
    MAX_DESC_CHARS = 42
    MAX_DETAIL_CHARS = 40
    
    # Calculer les totaux par taux de TVA
    total_ht_global = 0
    ht_par_taux = {}  # {taux: montant_ht}
    
    for i, ligne in enumerate(lignes_finales):
        ht_apres_remise = float(getattr(ligne, 'ht_apres_remise', 0) or 0)
        # IMPORTANT : 0 est une valeur valide pour TVA (auto-entrepreneur ou exonéré)
        # Ne pas utiliser "or tva_taux" car 0 serait remplacé par le taux global !
        ligne_tva = getattr(ligne, 'tva_taux', None)
        if ligne_tva is not None:
            tva_ligne = float(ligne_tva)
        else:
            tva_ligne = float(tva_taux)
        quantite = float(getattr(ligne, 'quantite', 1) or 1)
        unite = getattr(ligne, 'unite', 'u') or 'u'
        description = getattr(ligne, 'description', '') or ''
        description_detaillee = getattr(ligne, 'description_detaillee', '') or ''
        notes = getattr(ligne, 'notes', '') or ''
        
        # Le prix unitaire = HT après remise / quantité
        prix_unitaire = ht_apres_remise / quantite if quantite > 0 else ht_apres_remise
        
        total_ht_global += ht_apres_remise
        
        if tva_ligne not in ht_par_taux:
            ht_par_taux[tva_ligne] = 0
        ht_par_taux[tva_ligne] += ht_apres_remise
        
        print(f"   Ligne {i+1}: {description} | HT={ht_apres_remise:.2f}€ | TVA={tva_ligne}%")
        
        # Découper les textes en lignes
        lignes_desc_principale = decouper_texte_en_lignes(description, MAX_DESC_CHARS)
        lignes_desc_detaillee = decouper_texte_en_lignes(description_detaillee, MAX_DETAIL_CHARS)
        lignes_notes = decouper_texte_en_lignes(notes, MAX_DETAIL_CHARS - 6)
        
        # Calculer la hauteur de ligne
        nb_lignes_total = max(1, len(lignes_desc_principale))
        nb_lignes_total += len(lignes_desc_detaillee)
        nb_lignes_total += len(lignes_notes)
        
        if nb_lignes_total <= 1:
            hauteur_ligne = 10*mm
        else:
            hauteur_ligne = 8*mm + (nb_lignes_total * 3.5*mm)
        
        y_ligne -= hauteur_ligne
        
        # Fond alterné
        if i % 2 == 0:
            c.setFillColor(HexColor('#f8f9fa'))
            c.rect(15*mm, y_ligne, width - 30*mm, hauteur_ligne, fill=True, stroke=False)
        
        # Position Y pour le texte
        y_text = y_ligne + hauteur_ligne - 5*mm
        
        # Description principale
        c.setFillColor(GRIS_FONCE)
        c.setFont("Helvetica-Bold", 9)
        for j, ligne_desc in enumerate(lignes_desc_principale[:3]):
            c.drawString(18*mm, y_text, ligne_desc)
            y_text -= 3.5*mm
            if j == 0:
                c.setFont("Helvetica", 9)
        
        # Description détaillée
        if lignes_desc_detaillee:
            c.setFont("Helvetica", 7)
            c.setFillColor(HexColor('#555555'))
            for ligne_detail in lignes_desc_detaillee[:4]:
                c.drawString(18*mm, y_text, ligne_detail)
                y_text -= 3*mm
        
        # Notes
        if lignes_notes:
            c.setFont("Helvetica-Oblique", 7)
            c.setFillColor(HexColor('#777777'))
            for k, ligne_note in enumerate(lignes_notes[:2]):
                prefix = "Note: " if k == 0 else "      "
                c.drawString(18*mm, y_text, prefix + ligne_note)
                y_text -= 3*mm
        
        # Colonnes standard (alignées en haut)
        y_colonnes = y_ligne + hauteur_ligne - 5*mm
        c.setFont("Helvetica", 9)
        c.setFillColor(GRIS_FONCE)
        c.drawString(97*mm, y_colonnes, str(quantite))
        c.drawString(108*mm, y_colonnes, unite)
        c.drawString(125*mm, y_colonnes, f"{prix_unitaire:.2f} €")
        c.drawString(150*mm, y_colonnes, f"{tva_ligne}%")
        c.drawRightString(width - 18*mm, y_colonnes, f"{ht_apres_remise:.2f} €")
    
    y_ligne -= 5*mm
    
    # Ligne de séparation
    c.setStrokeColor(GRIS_CLAIR)
    c.setLineWidth(1)
    c.line(15*mm, y_ligne, width - 15*mm, y_ligne)
    
    # ============================================================
    # CALCUL DES TOTAUX AVEC TVA PAR TAUX
    # ============================================================
    
    # Calcul TVA par taux
    tva_par_taux = {}
    for taux, montant_ht in ht_par_taux.items():
        if taux > 0:
            tva_par_taux[taux] = montant_ht * (taux / 100)
    
    montant_tva_total = sum(tva_par_taux.values())
    total_ttc_avant_acompte = total_ht_global + montant_tva_total
    
    # Acompte TTC déjà versé
    total_acompte_ttc = float(acompte_ttc) if acompte_ttc else 0
    acompte_ref_texte = f" ({', '.join(acompte_refs)})" if acompte_refs else ""
    
    # Reste à payer
    reste_a_payer = total_ttc_avant_acompte - total_acompte_ttc
    
    print(f"📊 CALCULS FACTURE FINALE (depuis lignes_finales):")
    print(f"   Total HT (après remise): {total_ht_global:.2f} €")
    print(f"   TVA par taux: {tva_par_taux}")
    print(f"   Total TTC avant acompte: {total_ttc_avant_acompte:.2f} €")
    print(f"   Acompte TTC déjà versé: {total_acompte_ttc:.2f} €")
    print(f"   Reste à payer: {reste_a_payer:.2f} €")
    
    # ============================================================
    # AFFICHAGE DES TOTAUX AVEC REMISE
    # ============================================================
    y_totaux = y_ligne - 10*mm
    x_label = 130*mm
    x_value = width - 18*mm
    
    c.setFillColor(GRIS_FONCE)
    c.setFont("Helvetica", 10)
    
    y_offset = 0
    
    # Récupérer les informations de remise depuis data
    remise_type = getattr(data, 'remise_type', None)
    remise_valeur = getattr(data, 'remise_valeur', 0) or 0
    
    # Calculer le total HT avant remise si une remise est appliquée
    total_ht_avant_remise = total_ht_global
    remise_montant = 0
    
    if remise_type and remise_valeur > 0:
        if remise_type == "pourcentage":
            # total_ht_global = total_avant * (1 - remise/100)
            # donc total_avant = total_ht_global / (1 - remise/100)
            total_ht_avant_remise = total_ht_global / (1 - remise_valeur / 100)
            remise_montant = total_ht_avant_remise - total_ht_global
        elif remise_type in ["montant", "fixe"]:
            total_ht_avant_remise = total_ht_global + remise_valeur
            remise_montant = remise_valeur
    
    # Afficher Total HT avant remise (si remise présente)
    if remise_montant > 0:
        c.drawString(x_label, y_totaux - y_offset, "Total HT avant remise")
        c.drawRightString(x_value, y_totaux - y_offset, f"{total_ht_avant_remise:.2f} €")
        y_offset += 6*mm
        
        # Afficher la remise
        if remise_type == "pourcentage":
            c.drawString(x_label, y_totaux - y_offset, f"Remise ({remise_valeur}%)")
        else:
            c.drawString(x_label, y_totaux - y_offset, "Remise")
        c.setFillColor(HexColor('#e74c3c'))  # Rouge pour la remise
        c.drawRightString(x_value, y_totaux - y_offset, f"-{remise_montant:.2f} €")
        c.setFillColor(GRIS_FONCE)
        y_offset += 6*mm
        
        # Total HT après remise
        c.drawString(x_label, y_totaux - y_offset, "Total HT après remise")
        c.drawRightString(x_value, y_totaux - y_offset, f"{total_ht_global:.2f} €")
        y_offset += 6*mm
    else:
        # Pas de remise - Total HT simple
        c.drawString(x_label, y_totaux - y_offset, "Total HT")
        c.drawRightString(x_value, y_totaux - y_offset, f"{total_ht_global:.2f} €")
        y_offset += 6*mm
    
    # TVA par taux
    tva_affichee = False
    for taux in sorted(tva_par_taux.keys(), reverse=True):
        montant = tva_par_taux[taux]
        if taux > 0 and montant > 0:
            c.drawString(x_label, y_totaux - y_offset, f"TVA ({taux}%)")
            c.drawRightString(x_value, y_totaux - y_offset, f"{montant:.2f} €")
            y_offset += 6*mm
            tva_affichee = True
    
    # Si aucune TVA affichée (auto-entrepreneur)
    if not tva_affichee:
        c.setFont("Helvetica-Oblique", 9)
        c.drawString(x_label, y_totaux - y_offset, "TVA non applicable")
        c.setFont("Helvetica", 10)
        y_offset += 6*mm
    
    # Total TTC avant acompte (si acompte présent)
    if total_acompte_ttc > 0:
        c.setFont("Helvetica-Bold", 10)
        c.drawString(x_label, y_totaux - y_offset, "Total TTC")
        c.drawRightString(x_value, y_totaux - y_offset, f"{total_ttc_avant_acompte:.2f} €")
        y_offset += 8*mm
        
        # Acompte déjà versé - Ligne 1 : libellé + montant
        c.setFont("Helvetica", 10)
        c.setFillColor(GRIS_FONCE)
        c.drawString(x_label, y_totaux - y_offset, "Acompte déjà versé")
        c.setFillColor(HexColor('#27ae60'))  # Vert
        c.setFont("Helvetica-Bold", 10)
        c.drawRightString(x_value, y_totaux - y_offset, f"-{total_acompte_ttc:.2f} €")
        y_offset += 5*mm
        
        # Acompte - Ligne 2 : référence de la facture (en petit, italique)
        if acompte_refs and len(acompte_refs) > 0:
            c.setFont("Helvetica-Oblique", 8)
            c.setFillColor(HexColor('#666666'))
            refs_text = ", ".join(acompte_refs)
            c.drawString(x_label, y_totaux - y_offset, f"(Facture {refs_text})")
            y_offset += 6*mm
        else:
            y_offset += 3*mm
        
        c.setFillColor(GRIS_FONCE)
        
        # Vérifier si la facture est payée
        est_payee = getattr(data, 'statut', None) == 'payee'
        montant_reste_a_payer = 0.0 if est_payee else reste_a_payer
        
        # Encadré RESTE À PAYER (ou 0€ si payée)
        if est_payee:
            c.setFillColor(HexColor('#27ae60'))  # Vert pour payée
        else:
            c.setFillColor(get_couleur_principale(data))
        c.roundRect(x_label - 5*mm, y_totaux - y_offset - 8*mm, 68*mm, 10*mm, 2*mm, fill=True, stroke=False)
        c.setFillColor(white)
        c.setFont("Helvetica-Bold", 11)
        if est_payee:
            c.drawString(x_label, y_totaux - y_offset - 5*mm, "RESTE À PAYER")
            c.drawRightString(x_value, y_totaux - y_offset - 5*mm, "0,00 €")
        else:
            c.drawString(x_label, y_totaux - y_offset - 5*mm, "RESTE À PAYER")
            c.drawRightString(x_value, y_totaux - y_offset - 5*mm, f"{montant_reste_a_payer:.2f} €")
        
        return y_totaux - y_offset - 13*mm, total_ht_global, montant_reste_a_payer
    else:
        # Pas d'acompte - Total TTC simple
        est_payee = getattr(data, 'statut', None) == 'payee'
        
        if est_payee:
            c.setFillColor(HexColor('#27ae60'))  # Vert pour payée
        else:
            c.setFillColor(get_couleur_principale(data))
        c.roundRect(x_label - 5*mm, y_totaux - y_offset - 8*mm, 68*mm, 10*mm, 2*mm, fill=True, stroke=False)
        c.setFillColor(white)
        c.setFont("Helvetica-Bold", 11)
        if est_payee:
            c.drawString(x_label, y_totaux - y_offset - 5*mm, "RESTE À PAYER")
            c.drawRightString(x_value, y_totaux - y_offset - 5*mm, "0,00 €")
        else:
            c.drawString(x_label, y_totaux - y_offset - 5*mm, "TOTAL TTC")
            c.drawRightString(x_value, y_totaux - y_offset - 5*mm, f"{total_ttc_avant_acompte:.2f} €")
        
        return y_totaux - y_offset - 13*mm, total_ht_global, 0.0 if est_payee else total_ttc_avant_acompte


def dessiner_tableau_prestations(c, width, data, y_table, tva_taux):
    """Dessine le tableau des prestations pour une facture avec totaux propres"""
    
    # ============================================================
    # DÉTECTION DU TYPE DE FACTURE ET DES DONNÉES DISPONIBLES
    # ============================================================
    is_facture_acompte = getattr(data, 'is_facture_acompte', False)
    total_ttc_fourni = getattr(data, 'total_ttc', None)
    total_ht_fourni = getattr(data, 'total_ht', None)
    acompte_ttc_deja_facture = getattr(data, 'acompte_ttc_deja_facture', None)
    acompte_references = getattr(data, 'acompte_references', []) or []
    lignes_finales_devis = getattr(data, 'lignes_finales_devis', None)
    
    print(f"📄 FACTURE - is_facture_acompte: {is_facture_acompte}")
    print(f"   total_ttc_fourni: {total_ttc_fourni}, total_ht_fourni: {total_ht_fourni}")
    print(f"   acompte_ttc_deja_facture: {acompte_ttc_deja_facture}")
    print(f"   acompte_references: {acompte_references}")
    print(f"   lignes_finales_devis: {'OUI' if lignes_finales_devis and len(lignes_finales_devis) > 0 else 'NON'}")
    
    # ============================================================
    # PRIORITÉ : UTILISER lignes_finales_devis SI DISPONIBLE
    # ============================================================
    # Ces lignes contiennent les montants HT après remise et les TVA par ligne
    # C'est la source de vérité pour les factures finales
    
    if lignes_finales_devis and len(lignes_finales_devis) > 0:
        print(f"✅ UTILISATION DE lignes_finales_devis ({len(lignes_finales_devis)} lignes)")
        return dessiner_facture_depuis_lignes_finales(c, width, data, y_table, tva_taux, lignes_finales_devis, acompte_ttc_deja_facture, acompte_references)
    
    # ============================================================
    # FALLBACK : SÉPARER PRESTATIONS POSITIVES ET LIGNES D'ACOMPTE
    # ============================================================
    prestations_positives = []
    lignes_acompte = []
    
    for prestation in data.prestations:
        total_ligne = prestation.quantite * prestation.prix_unitaire
        desc = getattr(prestation, 'description', '').lower()
        
        # Si c'est une ligne d'acompte (prix négatif ou description contient "acompte")
        if total_ligne < 0 or 'acompte' in desc:
            lignes_acompte.append(prestation)
        else:
            prestations_positives.append(prestation)
    
    print(f"   Prestations positives: {len(prestations_positives)}, Lignes acompte: {len(lignes_acompte)}")
    
    # ============================================================
    # CAS FACTURE D'ACOMPTE : Affichage ventilé par taux de TVA
    # ============================================================
    if is_facture_acompte and total_ttc_fourni is not None:
        # En-tête du tableau
        c.setFillColor(get_couleur_principale(data))
        c.rect(15*mm, y_table, width - 30*mm, 10*mm, fill=True, stroke=False)
        
        c.setFillColor(white)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(18*mm, y_table + 3*mm, "Description")
        c.drawString(95*mm, y_table + 3*mm, "Qté")
        c.drawString(108*mm, y_table + 3*mm, "Unité")
        c.drawString(125*mm, y_table + 3*mm, "P.U. HT")
        c.drawString(150*mm, y_table + 3*mm, "TVA")
        c.drawRightString(width - 18*mm, y_table + 3*mm, "Total HT")
        
        y_ligne = y_table - 2*mm
        
        # Calculer les totaux par taux de TVA
        tva_par_taux = {}
        total_ht_calc = 0
        
        # Dessiner chaque prestation (ventilée par TVA)
        for idx, prestation in enumerate(data.prestations):
            y_ligne -= 10*mm
            
            # Alternance de couleur de fond
            if idx % 2 == 0:
                c.setFillColor(HexColor('#f8f9fa'))
            else:
                c.setFillColor(white)
            c.rect(15*mm, y_ligne, width - 30*mm, 10*mm, fill=True, stroke=False)
            
            c.setFillColor(GRIS_FONCE)
            c.setFont("Helvetica", 9)
            
            # Récupérer les valeurs
            desc = getattr(prestation, 'description', 'Acompte')
            quantite = float(getattr(prestation, 'quantite', 1) or 1)
            unite = getattr(prestation, 'unite', '') or ''
            prix_unitaire = float(getattr(prestation, 'prix_unitaire', 0) or 0)
            
            # Récupérer le taux TVA de la prestation
            presta_tva = getattr(prestation, 'tva_taux', None)
            if presta_tva is not None:
                tva_prestation = float(presta_tva)
            else:
                tva_prestation = tva_taux
            
            total_ht_ligne = quantite * prix_unitaire
            total_ht_calc += total_ht_ligne
            
            # Calculer et stocker la TVA
            montant_tva_ligne = total_ht_ligne * (tva_prestation / 100)
            if tva_prestation not in tva_par_taux:
                tva_par_taux[tva_prestation] = 0
            tva_par_taux[tva_prestation] += montant_tva_ligne
            
            # Dessiner la ligne
            c.drawString(18*mm, y_ligne + 2*mm, tronquer_texte(desc, 45))
            c.drawString(97*mm, y_ligne + 2*mm, str(int(quantite)) if quantite == int(quantite) else f"{quantite:.1f}")
            c.drawString(108*mm, y_ligne + 2*mm, unite)
            c.drawString(125*mm, y_ligne + 2*mm, f"{prix_unitaire:.2f} €")
            c.drawString(150*mm, y_ligne + 2*mm, f"{tva_prestation:.1f}%")
            c.drawRightString(width - 18*mm, y_ligne + 2*mm, f"{total_ht_ligne:.2f} €")
        
        # Ligne de séparation
        y_ligne -= 5*mm
        c.setStrokeColor(GRIS_CLAIR)
        c.setLineWidth(1)
        c.line(15*mm, y_ligne, width - 15*mm, y_ligne)
        
        # Calculer le total TVA
        total_tva_calc = sum(tva_par_taux.values())
        total_ttc_calc = total_ht_calc + total_tva_calc
        
        # Utiliser les valeurs fournies si disponibles
        total_ttc = float(total_ttc_fourni)
        total_ht_final = float(total_ht_fourni) if total_ht_fourni is not None else total_ht_calc
        
        # Totaux
        y_totaux = y_ligne - 10*mm
        x_label = 130*mm
        x_value = width - 18*mm
        
        c.setFillColor(GRIS_FONCE)
        c.setFont("Helvetica", 10)
        c.drawString(x_label, y_totaux, "Total HT")
        c.drawRightString(x_value, y_totaux, f"{total_ht_final:.2f} €")
        
        y_offset = 6*mm
        
        # Afficher la TVA par taux
        tva_affichee = False
        for taux_tva, montant_tva in sorted(tva_par_taux.items()):
            if montant_tva > 0.01:
                c.drawString(x_label, y_totaux - y_offset, f"TVA ({taux_tva:.1f}%)")
                c.drawRightString(x_value, y_totaux - y_offset, f"{montant_tva:.2f} €")
                y_offset += 6*mm
                tva_affichee = True
        
        # Si aucune TVA (toutes à 0%), afficher "TVA non applicable"
        if not tva_affichee:
            c.drawString(x_label, y_totaux - y_offset, "TVA non applicable")
            y_offset += 6*mm
        
        # Total TTC avec fond coloré
        y_offset += 2*mm
        c.setFillColor(get_couleur_principale(data))
        c.rect(x_label - 5*mm, y_totaux - y_offset - 3*mm, width - x_label - 5*mm, 10*mm, fill=True, stroke=False)
        c.setFillColor(white)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(x_label, y_totaux - y_offset, "TOTAL TTC")
        c.drawRightString(x_value, y_totaux - y_offset, f"{total_ttc:.2f} €")
        
        return y_totaux - y_offset - 8*mm, total_ht_final, total_ttc
    
    # ============================================================
    # CAS FACTURE FINALE/NORMALE : Calcul complet avec TVA par taux
    # ============================================================
    
    # En-tête du tableau
    c.setFillColor(get_couleur_principale(data))
    c.rect(15*mm, y_table, width - 30*mm, 10*mm, fill=True, stroke=False)
    
    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(18*mm, y_table + 3*mm, "Description")
    c.drawString(95*mm, y_table + 3*mm, "Qté")
    c.drawString(108*mm, y_table + 3*mm, "Unité")
    c.drawString(125*mm, y_table + 3*mm, "P.U. HT")
    c.drawString(150*mm, y_table + 3*mm, "TVA")
    c.drawRightString(width - 18*mm, y_table + 3*mm, "Total HT")
    
    y_ligne = y_table - 2*mm
    
    # Largeur max pour les descriptions
    MAX_DESC_CHARS = 42
    MAX_DETAIL_CHARS = 40
    
    # Calcul des totaux HT et TVA par taux (seulement prestations positives)
    total_ht_avant_remise = 0
    ht_par_taux = {}  # {taux: montant_ht}
    
    for i, prestation in enumerate(prestations_positives):
        total_ligne = prestation.quantite * prestation.prix_unitaire
        total_ht_avant_remise += total_ligne
        
        # Récupérer le taux TVA de la prestation
        tva_prestation_raw = getattr(prestation, 'tva_taux', None)
        tva_prestation = tva_prestation_raw if tva_prestation_raw is not None else tva_taux
        
        if tva_prestation not in ht_par_taux:
            ht_par_taux[tva_prestation] = 0
        ht_par_taux[tva_prestation] += total_ligne
        
        # Récupérer les textes
        description_principale = getattr(prestation, 'description', '') or ''
        description_detaillee = getattr(prestation, 'description_detaillee', '') or ''
        notes = getattr(prestation, 'notes', '') or ''
        
        # Découper les textes en lignes
        lignes_desc_principale = decouper_texte_en_lignes(description_principale, MAX_DESC_CHARS)
        lignes_desc_detaillee = decouper_texte_en_lignes(description_detaillee, MAX_DETAIL_CHARS)
        lignes_notes = decouper_texte_en_lignes(notes, MAX_DETAIL_CHARS - 6)
        
        # Calculer la hauteur de ligne
        nb_lignes_total = max(1, len(lignes_desc_principale))
        nb_lignes_total += len(lignes_desc_detaillee)
        nb_lignes_total += len(lignes_notes)
        
        if nb_lignes_total <= 1:
            hauteur_ligne = 10*mm
        else:
            hauteur_ligne = 8*mm + (nb_lignes_total * 3.5*mm)
        
        y_ligne -= hauteur_ligne
        
        # Fond alterné
        if i % 2 == 0:
            c.setFillColor(HexColor('#f8f9fa'))
            c.rect(15*mm, y_ligne, width - 30*mm, hauteur_ligne, fill=True, stroke=False)
        
        # Position Y pour le texte
        y_text = y_ligne + hauteur_ligne - 5*mm
        
        # Description principale
        c.setFillColor(GRIS_FONCE)
        c.setFont("Helvetica-Bold", 9)
        for j, ligne_desc in enumerate(lignes_desc_principale[:3]):
            c.drawString(18*mm, y_text, ligne_desc)
            y_text -= 3.5*mm
            if j == 0:
                c.setFont("Helvetica", 9)
        
        # Description détaillée
        if lignes_desc_detaillee:
            c.setFont("Helvetica", 7)
            c.setFillColor(HexColor('#555555'))
            for ligne_detail in lignes_desc_detaillee[:4]:
                c.drawString(18*mm, y_text, ligne_detail)
                y_text -= 3*mm
        
        # Notes
        if lignes_notes:
            c.setFont("Helvetica-Oblique", 7)
            c.setFillColor(HexColor('#777777'))
            for k, ligne_note in enumerate(lignes_notes[:2]):
                prefix = "Note: " if k == 0 else "      "
                c.drawString(18*mm, y_text, prefix + ligne_note)
                y_text -= 3*mm
        
        # Colonnes standard (alignées en haut)
        y_colonnes = y_ligne + hauteur_ligne - 5*mm
        c.setFont("Helvetica", 9)
        c.setFillColor(GRIS_FONCE)
        c.drawString(97*mm, y_colonnes, str(prestation.quantite))
        c.drawString(108*mm, y_colonnes, getattr(prestation, 'unite', 'u') or 'u')
        c.drawString(125*mm, y_colonnes, f"{prestation.prix_unitaire:.2f} €")
        c.drawString(150*mm, y_colonnes, f"{tva_prestation}%")
        c.drawRightString(width - 18*mm, y_colonnes, f"{total_ligne:.2f} €")
    
    y_ligne -= 5*mm
    
    # Ligne de séparation
    c.setStrokeColor(GRIS_CLAIR)
    c.setLineWidth(1)
    c.line(15*mm, y_ligne, width - 15*mm, y_ligne)
    
    # ============================================================
    # CALCUL DES TOTAUX AVEC REMISE ET TVA PAR TAUX
    # ============================================================
    
    # Calcul de la remise
    remise = 0
    remise_type = getattr(data, 'remise_type', None)
    remise_valeur = getattr(data, 'remise_valeur', 0) or 0
    
    if remise_type and remise_valeur > 0:
        if remise_type == "pourcentage":
            remise = total_ht_avant_remise * (remise_valeur / 100)
        elif remise_type in ["montant", "fixe"]:
            remise = remise_valeur
    
    total_ht_apres_remise = total_ht_avant_remise - remise
    
    # Ratio remise pour calculer HT par taux après remise
    ratio_remise = total_ht_apres_remise / total_ht_avant_remise if total_ht_avant_remise > 0 else 1
    
    # Calcul TVA par taux (après remise)
    tva_par_taux = {}
    for taux, montant_ht in ht_par_taux.items():
        montant_ht_apres_remise = montant_ht * ratio_remise
        if taux > 0:
            tva_par_taux[taux] = montant_ht_apres_remise * (taux / 100)
    
    montant_tva_total = sum(tva_par_taux.values())
    total_ttc_avant_acompte = total_ht_apres_remise + montant_tva_total
    
    # Calcul de l'acompte à déduire
    total_acompte_ttc = 0
    acompte_ref_texte = ""
    
    # 1. Depuis acompte_ttc_deja_facture (envoyé par le frontend)
    if acompte_ttc_deja_facture and float(acompte_ttc_deja_facture) > 0:
        total_acompte_ttc = float(acompte_ttc_deja_facture)
        if acompte_references:
            acompte_ref_texte = f" ({', '.join(acompte_references)})"
    
    # 2. Sinon, depuis les lignes d'acompte négatives
    elif lignes_acompte:
        for ligne in lignes_acompte:
            total_acompte_ttc += abs(ligne.quantite * ligne.prix_unitaire)
        acompte_ref_texte = ""
    
    # Reste à payer
    reste_a_payer = total_ttc_avant_acompte - total_acompte_ttc
    
    print(f"📊 CALCULS FACTURE FINALE:")
    print(f"   Total HT avant remise: {total_ht_avant_remise:.2f} €")
    print(f"   Remise ({remise_type}): {remise:.2f} €")
    print(f"   Total HT après remise: {total_ht_apres_remise:.2f} €")
    print(f"   TVA par taux: {tva_par_taux}")
    print(f"   Total TTC avant acompte: {total_ttc_avant_acompte:.2f} €")
    print(f"   Acompte TTC déjà versé: {total_acompte_ttc:.2f} €")
    print(f"   Reste à payer: {reste_a_payer:.2f} €")
    
    # ============================================================
    # AFFICHAGE DES TOTAUX
    # ============================================================
    y_totaux = y_ligne - 10*mm
    x_label = 130*mm
    x_value = width - 18*mm
    
    c.setFillColor(GRIS_FONCE)
    c.setFont("Helvetica", 10)
    
    y_offset = 0
    
    # Total HT avant remise (ou Total HT si pas de remise)
    if remise > 0:
        c.drawString(x_label, y_totaux - y_offset, "Total HT avant remise")
    else:
        c.drawString(x_label, y_totaux - y_offset, "Total HT")
    c.drawRightString(x_value, y_totaux - y_offset, f"{total_ht_avant_remise:.2f} €")
    y_offset += 6*mm
    
    # Remise si présente
    if remise > 0:
        if remise_type == "pourcentage":
            c.drawString(x_label, y_totaux - y_offset, f"Remise ({remise_valeur}%)")
        else:
            c.drawString(x_label, y_totaux - y_offset, "Remise")
        c.setFillColor(HexColor('#e74c3c'))
        c.drawRightString(x_value, y_totaux - y_offset, f"-{remise:.2f} €")
        c.setFillColor(GRIS_FONCE)
        y_offset += 6*mm
    
        # Total HT après remise
        c.drawString(x_label, y_totaux - y_offset, "Total HT après remise")
        c.drawRightString(x_value, y_totaux - y_offset, f"{total_ht_apres_remise:.2f} €")
        y_offset += 6*mm
    
    # TVA par taux
    tva_affichee = False
    for taux in sorted(tva_par_taux.keys(), reverse=True):
        montant = tva_par_taux[taux]
        if taux > 0 and montant > 0:
            c.drawString(x_label, y_totaux - y_offset, f"TVA ({taux}%)")
            c.drawRightString(x_value, y_totaux - y_offset, f"{montant:.2f} €")
            y_offset += 6*mm
            tva_affichee = True
    
    # Si aucune TVA affichée (auto-entrepreneur)
    if not tva_affichee:
        c.setFont("Helvetica-Oblique", 9)
        c.drawString(x_label, y_totaux - y_offset, "TVA non applicable")
        c.setFont("Helvetica", 10)
        y_offset += 6*mm
    
    # Total TTC avant acompte (si acompte présent)
    if total_acompte_ttc > 0:
        c.setFont("Helvetica-Bold", 10)
        c.drawString(x_label, y_totaux - y_offset, "Total TTC")
        c.drawRightString(x_value, y_totaux - y_offset, f"{total_ttc_avant_acompte:.2f} €")
        y_offset += 8*mm
        
        # Acompte déjà versé - Ligne 1 : libellé + montant
        c.setFont("Helvetica", 10)
        c.setFillColor(GRIS_FONCE)
        c.drawString(x_label, y_totaux - y_offset, "Acompte déjà versé")
        c.setFillColor(HexColor('#27ae60'))  # Vert
        c.setFont("Helvetica-Bold", 10)
        c.drawRightString(x_value, y_totaux - y_offset, f"-{total_acompte_ttc:.2f} €")
        y_offset += 5*mm
        
        # Acompte - Ligne 2 : référence de la facture (en petit, italique)
        if acompte_references and len(acompte_references) > 0:
            c.setFont("Helvetica-Oblique", 8)
            c.setFillColor(HexColor('#666666'))
            refs_text = ", ".join(acompte_references)
            c.drawString(x_label, y_totaux - y_offset, f"(Facture {refs_text})")
            y_offset += 6*mm
        elif acompte_ref_texte:
            # Fallback pour l'ancien format
            c.setFont("Helvetica-Oblique", 8)
            c.setFillColor(HexColor('#666666'))
            c.drawString(x_label, y_totaux - y_offset, acompte_ref_texte.strip())
            y_offset += 6*mm
        else:
            y_offset += 3*mm
        
        c.setFillColor(GRIS_FONCE)
        
        # Vérifier si la facture est payée
        est_payee = getattr(data, 'statut', None) == 'payee'
        montant_reste_a_payer = 0.0 if est_payee else reste_a_payer
        
        # Encadré RESTE À PAYER (ou PAYÉ si statut payee)
        if est_payee:
            c.setFillColor(HexColor('#27ae60'))  # Vert pour payée
        else:
            c.setFillColor(get_couleur_principale(data))
        c.roundRect(x_label - 5*mm, y_totaux - y_offset - 8*mm, 68*mm, 10*mm, 2*mm, fill=True, stroke=False)
        c.setFillColor(white)
        c.setFont("Helvetica-Bold", 11)
        if est_payee:
            c.drawString(x_label, y_totaux - y_offset - 5*mm, "RESTE À PAYER")
            c.drawRightString(x_value, y_totaux - y_offset - 5*mm, "0,00 €")
        else:
            c.drawString(x_label, y_totaux - y_offset - 5*mm, "RESTE À PAYER")
            c.drawRightString(x_value, y_totaux - y_offset - 5*mm, f"{montant_reste_a_payer:.2f} €")
        
        return y_totaux - y_offset - 13*mm, total_ht_apres_remise, montant_reste_a_payer
    else:
        # Pas d'acompte - Total TTC simple
        est_payee = getattr(data, 'statut', None) == 'payee'
        
        if est_payee:
            c.setFillColor(HexColor('#27ae60'))  # Vert pour payée
        else:
            c.setFillColor(get_couleur_principale(data))
        c.roundRect(x_label - 5*mm, y_totaux - y_offset - 8*mm, 68*mm, 10*mm, 2*mm, fill=True, stroke=False)
        c.setFillColor(white)
        c.setFont("Helvetica-Bold", 11)
        if est_payee:
            c.drawString(x_label, y_totaux - y_offset - 5*mm, "RESTE À PAYER")
            c.drawRightString(x_value, y_totaux - y_offset - 5*mm, "0,00 €")
        else:
            c.drawString(x_label, y_totaux - y_offset - 5*mm, "TOTAL TTC")
            c.drawRightString(x_value, y_totaux - y_offset - 5*mm, f"{total_ttc_avant_acompte:.2f} €")
        
        return y_totaux - y_offset - 13*mm, total_ht_apres_remise, 0.0 if est_payee else total_ttc_avant_acompte


def dessiner_pied_page(c, width, data, mention_tva=""):
    c.setStrokeColor(get_couleur_principale(data))
    c.setLineWidth(2)
    c.line(15*mm, 35*mm, width - 15*mm, 35*mm)
    
    c.setFillColor(GRIS_TEXTE)
    c.setFont("Helvetica", 7)
    
    # Récupérer les infos de forme juridique
    forme_raw = getattr(data.entreprise, 'forme_juridique', None)
    forme = forme_raw.lower().strip() if forme_raw and forme_raw.strip() else None
    capital = getattr(data.entreprise, 'capital_social', '') or ''
    rcs = getattr(data.entreprise, 'rcs', '') or ''
    tva_intra = getattr(data.entreprise, 'tva_intracommunautaire', '') or ''
    
    # Ligne 1 : Nom + forme juridique + capital (si applicable)
    if forme in ['sarl', 'eurl', 'sas', 'sasu']:
        ligne1 = f"{data.entreprise.nom} - {forme.upper()}"
        if capital:
            ligne1 += f" au capital de {capital} €"
    elif forme in ['ei']:
        ligne1 = f"{data.entreprise.nom} - Entreprise Individuelle"
    elif forme in ['auto-entrepreneur', 'micro-entreprise', 'autoentrepreneur', 'microentreprise']:
        ligne1 = f"{data.entreprise.nom} - Auto-entrepreneur"
    else:
        # Si pas de forme juridique définie, juste le nom
        ligne1 = f"{data.entreprise.nom}"
    
    c.drawCentredString(width/2, 28*mm, ligne1)
    
    # Ligne 2 : SIRET + RCS (si applicable)
    ligne2 = f"SIRET : {data.entreprise.siret}"
    if rcs and forme in ['sarl', 'eurl', 'sas', 'sasu']:
        ligne2 += f" - {rcs}"
    elif forme in ['auto-entrepreneur', 'micro-entreprise', 'autoentrepreneur', 'microentreprise']:
        ligne2 += " - Dispensé d'immatriculation au RCS"
    
    c.drawCentredString(width/2, 23*mm, ligne2)
    
    # Ligne 3 : Adresse + Tél
    adresse_pied = formater_adresse_complete(data.entreprise.adresse, data.entreprise.cp_ville)
    c.drawCentredString(width/2, 18*mm, f"{adresse_pied} - Tél : {data.entreprise.tel}")
    
    # Ligne 4 : TVA
    if mention_tva:
        c.setFont("Helvetica-Oblique", 7)
        c.drawCentredString(width/2, 13*mm, mention_tva)
    elif tva_intra:
        c.drawCentredString(width/2, 13*mm, f"N° TVA intracommunautaire : {tva_intra}")
    else:
        siret_clean = data.entreprise.siret.replace(' ', '').replace('.', '')
        c.drawCentredString(width/2, 13*mm, f"TVA intracommunautaire : FR{siret_clean[:9] if len(siret_clean) >= 9 else siret_clean}")
    
    c.setFillColor(get_couleur_principale(data))
    c.setFont("Helvetica-Oblique", 6)
    c.drawRightString(width - 15*mm, 8*mm, "Généré par Vocario.fr")


def generer_pdf_devis(data: DevisRequest, numero_devis_force: Optional[str] = None) -> str:
    # PRIORITÉ 1: Utiliser le numéro forcé (paramètre explicite)
    # PRIORITÉ 2: Utiliser le numéro fourni dans data.numero_devis
    # PRIORITÉ 3: Générer un nouveau numéro (ne devrait jamais arriver)
    
    if numero_devis_force and str(numero_devis_force).strip():
        numero_devis = str(numero_devis_force).strip()
        print(f"✅ Utilisation du numéro de devis FORCÉ (paramètre): '{numero_devis}'")
    elif data.numero_devis and str(data.numero_devis).strip():
        numero_devis = str(data.numero_devis).strip()
        print(f"✅ Utilisation du numéro de devis fourni dans data: '{numero_devis}'")
    else:
        # Si aucun numéro n'est fourni, c'est une erreur critique
        numero_devis = f"DEV-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"
        print(f"❌ ERREUR CRITIQUE: numero_devis non fourni ou vide!")
        print(f"   - numero_devis_force = '{numero_devis_force}'")
        print(f"   - data.numero_devis = '{data.numero_devis}'")
        print(f"   - Génération d'un nouveau numéro (ce ne devrait pas arriver): {numero_devis}")
        print(f"⚠️ ATTENTION: Le numéro généré ({numero_devis}) ne correspondra pas au numéro en base de données!")
    
    filename = f"{numero_devis}.pdf"
    filepath = os.path.join(PDF_FOLDER, filename)
    
    date_validite = (datetime.now() + timedelta(days=data.validite_jours)).strftime("%d/%m/%Y")
    
    logo = telecharger_logo(data.entreprise.logo_url)
    
    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4
    
    # Calculer les totaux globaux sur toutes les prestations
    total_ht_avant_acompte = 0
    total_acompte = 0
    for prestation in data.prestations:
        total_ligne = prestation.quantite * prestation.prix_unitaire
        if total_ligne >= 0:
            total_ht_avant_acompte += total_ligne
        else:
            total_acompte += abs(total_ligne)
    
    # Calcul de la remise directement à partir de data.remise_type et data.remise_valeur
    remise_type = getattr(data, 'remise_type', None)
    remise_valeur = getattr(data, 'remise_valeur', 0) or 0
    
    # Normaliser remise_type
    if remise_type:
        remise_type = str(remise_type).strip()
        if remise_type == "" or remise_type.lower() == "none":
            remise_type = None
    
    # Convertir remise_valeur en nombre
    try:
        remise_valeur = float(remise_valeur)
    except (ValueError, TypeError):
        remise_valeur = 0
    
    # Calculer remise_totale à partir de remise_type et remise_valeur
    if remise_type == "pourcentage" and remise_valeur > 0:
        remise = total_ht_avant_acompte * (remise_valeur / 100)
        print(f"✅ Remise pourcentage calculée: {remise:.2f} € ({remise_valeur}% de {total_ht_avant_acompte:.2f})")
    elif remise_type in ["montant", "fixe"] and remise_valeur > 0:
        remise = remise_valeur
        print(f"✅ Remise montant calculée: {remise:.2f} €")
    else:
        remise = 0
        if remise_type:
            print(f"⚠️ Remise_type défini ('{remise_type}') mais remise_valeur invalide: {remise_valeur}")
        else:
            print(f"ℹ️ Pas de remise définie")
    
    # Appliquer la remise, puis déduire l'acompte
    total_ht_apres_remise = total_ht_avant_acompte - remise
    total_ht_final = total_ht_apres_remise - total_acompte
    montant_tva = total_ht_final * (data.tva_taux / 100)
    total_ttc = total_ht_final + montant_tva
    total_ht = total_ht_avant_acompte  # Pour l'affichage
    
    # Stocker la remise dans data pour qu'elle soit accessible dans dessiner_totaux
    # On utilise une approche différente : on va passer les valeurs directement
    print(f"📋 Données finales - remise: {remise:.2f}, remise_type dans data: '{getattr(data, 'remise_type', None)}', remise_valeur dans data: {getattr(data, 'remise_valeur', None)}")
    
    # Pagination : diviser les prestations en groupes
    lignes_par_page = 11  # Nombre de lignes par page
    prestations_groupes = []
    for i in range(0, len(data.prestations), lignes_par_page):
        prestations_groupes.append(data.prestations[i:i + lignes_par_page])
    
    # Si aucune prestation, créer au moins une page vide
    if not prestations_groupes:
        prestations_groupes = [[]]
    
    mention_tva = ""
    if data.tva_taux == 0:
        mention_tva = "TVA non applicable, article 293 B du Code général des impôts"
    
    # Dessiner chaque groupe de prestations
    for page_num, groupe_prestations in enumerate(prestations_groupes):
        est_premiere_page = (page_num == 0)
        est_derniere_page = (page_num == len(prestations_groupes) - 1)
        
        # Dessiner l'en-tête de page
        dessiner_en_tete_page(c, width, height, data, numero_devis, logo, date_validite)
        
        if est_premiere_page:
            # Dessiner les blocs emetteur/client sur la première page uniquement
            y_position = height - 60*mm
            dessiner_bloc_emetteur(c, width, height, data, y_position)
            dessiner_bloc_client(c, width, height, data, y_position)
            
            c.setFillColor(GRIS_TEXTE)
            c.setFont("Helvetica", 9)
            c.drawRightString(width - 20*mm, y_position - 28*mm, f"Validité : {date_validite}")
            
            y_table = y_position - 50*mm
        else:
            # Sur les pages suivantes, le tableau commence plus haut
            y_table = height - 55*mm
        
        # Dessiner les lignes de prestations
        index_debut = page_num * lignes_par_page
        y_totaux_tableau, _, _ = dessiner_lignes_prestations(c, width, groupe_prestations, y_table, data, index_debut)
        
        # Si dernière page, dessiner les totaux, signature et conditions
        if est_derniere_page:
            y_totaux = y_totaux_tableau
            
            # Log avant dessiner_totaux pour vérifier les valeurs
            print(f"📊 AVANT dessiner_totaux - remise: {remise:.2f}, remise_type: '{getattr(data, 'remise_type', None)}', remise_valeur: {getattr(data, 'remise_valeur', None)}")
            
            # Dessiner les totaux
            y_fin_totaux = dessiner_totaux(c, width, y_totaux, total_ht, total_ht_avant_acompte, total_acompte, remise, data.tva_taux, total_ht_final, total_ttc, data)
            
            # Bloc signature À GAUCHE (au niveau des totaux)
            y_signature = y_totaux - 5*mm
            c.setStrokeColor(GRIS_CLAIR)
            c.setLineWidth(1)
            c.roundRect(15*mm, y_signature - 35*mm, 80*mm, 40*mm, 3*mm, fill=False, stroke=True)
            
            c.setFillColor(GRIS_TEXTE)
            c.setFont("Helvetica-Bold", 9)
            c.drawString(20*mm, y_signature - 3*mm, "Bon pour accord")
            c.setFont("Helvetica", 8)
            c.drawString(20*mm, y_signature - 13*mm, "Date :")
            c.drawString(20*mm, y_signature - 23*mm, "Signature :")
            c.setFont("Helvetica-Oblique", 7)
            c.drawString(20*mm, y_signature - 31*mm, "(Précédée de \"Bon pour accord\")")
            
            # Vérifier s'il y a assez d'espace pour les conditions APRÈS les totaux/signature
            hauteur_conditions = 35*mm
            espace_necessaire_conditions = hauteur_conditions + 40*mm  # 40mm marge pour le footer
            # Position des conditions après la signature (prendre le plus bas entre signature et totaux)
            y_bas_signature = y_signature - 35*mm
            y_conditions_possible = min(y_fin_totaux, y_bas_signature) - 45*mm
            
            # Si pas assez d'espace pour les conditions sur cette page, créer une nouvelle page
            if y_conditions_possible < espace_necessaire_conditions:
                # Dessiner le footer sur la page actuelle (avec totaux et signature)
                dessiner_pied_page(c, width, data, mention_tva)
                # Créer une nouvelle page pour les conditions
                c.showPage()
                dessiner_en_tete_page(c, width, height, data, numero_devis, logo, date_validite)
                y_conditions = height - 55*mm
            else:
                # Dessiner les conditions sur la même page, APRÈS les totaux/signature
                y_conditions = y_conditions_possible
            
            # Dessiner les conditions
            c.setFillColor(GRIS_CLAIR)
            c.roundRect(15*mm, y_conditions - 25*mm, width - 30*mm, 35*mm, 3*mm, fill=True, stroke=False)
            
            c.setFillColor(get_couleur_principale(data))
            c.setFont("Helvetica-Bold", 10)
            c.drawString(20*mm, y_conditions + 2*mm, "CONDITIONS")
            
            c.setFillColor(GRIS_FONCE)
            c.setFont("Helvetica", 9)
            c.drawString(20*mm, y_conditions - 8*mm, f"• Délai de réalisation : {data.delai_realisation}")
            c.drawString(20*mm, y_conditions - 14*mm, f"• Conditions de paiement : {data.entreprise.conditions_paiement or data.conditions_paiement}")
            c.drawString(20*mm, y_conditions - 20*mm, f"• Devis valable jusqu'au : {date_validite}")
            
            # Dessiner le footer sur cette page (avec totaux, signature et conditions)
            dessiner_pied_page(c, width, data, mention_tva)
        
        # Dessiner le footer sur chaque page (sauf la dernière page qui l'a déjà dessiné)
        if not est_derniere_page:
            dessiner_pied_page(c, width, data, mention_tva)
        
        # Si ce n'est pas la dernière page, créer une nouvelle page
        if not est_derniere_page:
            c.showPage()
    
    try:
        c.save()
        print(f"✅ PDF devis sauvegardé: {filepath}")
    except Exception as e:
        print(f"❌ Erreur lors de la sauvegarde du PDF: {e}")
        raise
    
    return filepath, numero_devis, total_ht_final, total_ttc


def generer_pdf_facture(data: FactureRequest, numero_facture_force: Optional[str] = None) -> str:
    # PRIORITÉ 1: Utiliser le numéro forcé (paramètre explicite)
    # PRIORITÉ 2: Utiliser le numéro fourni dans data.numero_facture
    # PRIORITÉ 3: Générer un nouveau numéro (ne devrait jamais arriver)
    
    if numero_facture_force and str(numero_facture_force).strip():
        numero_facture = str(numero_facture_force).strip()
        print(f"✅ Facture PDF - Utilisation du numéro FORCÉ (paramètre): '{numero_facture}'")
    elif data.numero_facture and str(data.numero_facture).strip():
        numero_facture = str(data.numero_facture).strip()
        print(f"✅ Facture PDF - Utilisation du numéro fourni dans data: '{numero_facture}'")
    else:
        # Si aucun numéro n'est fourni, c'est une erreur critique
        numero_facture = f"FAC-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"
        print(f"❌ ERREUR CRITIQUE: numero_facture non fourni ou vide!")
        print(f"   - numero_facture_force = '{numero_facture_force}'")
        print(f"   - data.numero_facture = '{data.numero_facture}'")
        print(f"   - Génération d'un nouveau numéro (ce ne devrait pas arriver): {numero_facture}")
        print(f"⚠️ ATTENTION: Le numéro généré ({numero_facture}) ne correspondra pas au numéro en base de données!")
    
    filename = f"{numero_facture}.pdf"
    filepath = os.path.join(PDF_FOLDER, filename)
    
    date_echeance = (datetime.now() + timedelta(days=data.date_echeance_jours)).strftime("%d/%m/%Y")
    
    logo = telecharger_logo(data.entreprise.logo_url)
    
    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4
    
    c.setFillColor(get_couleur_principale(data))
    c.rect(0, height - 45*mm, width, 45*mm, fill=True, stroke=False)
    
    text_start_x = 15*mm
    
    if logo:
        try:
            logo_size = 30*mm
            c.drawImage(logo, 15*mm, height - 40*mm, width=logo_size, height=logo_size, preserveAspectRatio=True, mask='auto')
            text_start_x = 50*mm
        except Exception as e:
            print(f"Erreur logo: {e}")
    
    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 18)
    c.drawString(text_start_x, height - 18*mm, tronquer_texte(data.entreprise.nom.upper(), 30))
    
    if data.entreprise.gerant and data.entreprise.gerant.strip():
        c.setFont("Helvetica", 9)
        c.drawString(text_start_x, height - 26*mm, f"Gérant : {data.entreprise.gerant}")
    
    c.setFont("Helvetica-Bold", 28)
    c.drawRightString(width - 20*mm, height - 18*mm, "FACTURE")
    c.setFont("Helvetica", 11)
    c.drawRightString(width - 20*mm, height - 28*mm, f"N° {numero_facture}")
    
    # Vérifier si la facture est payée
    est_payee = hasattr(data, 'statut') and data.statut == 'payee'
    
    if est_payee:
        # Afficher "PAYÉE" en vert à côté du numéro
        c.setFillColor(HexColor('#27ae60'))  # Vert pour "PAYÉE"
        c.setFont("Helvetica-Bold", 12)
        c.drawRightString(width - 20*mm, height - 36*mm, "PAYÉE")
        c.setFillColor(white)  # Remettre la couleur blanche pour la suite
    
    c.setFont("Helvetica", 9)
    c.setFillColor(white)
    y_date = height - 42*mm if est_payee else height - 36*mm
    c.drawRightString(width - 20*mm, y_date, f"Date : {datetime.now().strftime('%d/%m/%Y')}")
    
    if data.numero_devis_origine:
        c.setFont("Helvetica", 8)
        y_ref_devis = y_date - 6*mm
        c.drawRightString(width - 20*mm, y_ref_devis, f"Réf. devis : {data.numero_devis_origine}")
    
    y_position = height - 60*mm
    dessiner_bloc_emetteur(c, width, height, data, y_position)
    dessiner_bloc_client(c, width, height, data, y_position)
    
    c.setFillColor(GRIS_TEXTE)
    c.setFont("Helvetica", 9)
    if not est_payee:
        c.drawRightString(width - 20*mm, y_position - 28*mm, f"Échéance : {date_echeance}")
    
    y_table = y_position - 50*mm
    y_totaux, total_ht, total_ttc = dessiner_tableau_prestations(c, width, data, y_table, data.tva_taux)
    
    y_paiement = y_totaux - 45*mm
    c.setFillColor(GRIS_CLAIR)
    c.roundRect(15*mm, y_paiement - 30*mm, width - 30*mm, 40*mm, 3*mm, fill=True, stroke=False)
    
    c.setFillColor(get_couleur_principale(data))
    c.setFont("Helvetica-Bold", 10)
    c.drawString(20*mm, y_paiement + 2*mm, "INFORMATIONS DE PAIEMENT")
    
    c.setFillColor(GRIS_FONCE)
    c.setFont("Helvetica", 9)
    
    if est_payee:
        # Si la facture est payée, afficher "Reste à payer : 0 €"
        c.drawString(20*mm, y_paiement - 8*mm, f"• Reste à payer : 0,00 €")
        c.drawString(20*mm, y_paiement - 14*mm, "• Paiement reçu")
    else:
        # Sinon, afficher les informations normales
        c.drawString(20*mm, y_paiement - 8*mm, f"• Date d'échéance : {date_echeance}")
        c.drawString(20*mm, y_paiement - 14*mm, "• Mode de paiement : Virement bancaire, chèque ou espèces")
        c.drawString(20*mm, y_paiement - 20*mm, "• En cas de retard : pénalité de 3 fois le taux d'intérêt légal")
        c.drawString(20*mm, y_paiement - 26*mm, "• Indemnité forfaitaire pour frais de recouvrement : 40€")
    
    # Afficher le RIB si disponible
    if data.rib and data.rib.iban:
        y_rib = y_paiement - 45*mm
        c.setFillColor(GRIS_CLAIR)
        c.roundRect(15*mm, y_rib - 20*mm, width - 30*mm, 30*mm, 3*mm, fill=True, stroke=False)
        
        c.setFillColor(get_couleur_principale(data))
        c.setFont("Helvetica-Bold", 10)
        c.drawString(20*mm, y_rib + 2*mm, "COORDONNÉES BANCAIRES")
        
        c.setFillColor(GRIS_FONCE)
        c.setFont("Helvetica", 9)
        c.drawString(20*mm, y_rib - 6*mm, f"IBAN : {data.rib.iban}")
        c.drawString(20*mm, y_rib - 12*mm, f"BIC : {data.rib.bic}")
        if data.rib.titulaire:
            c.drawString(20*mm, y_rib - 18*mm, f"Titulaire : {data.rib.titulaire}")
    
    mention_tva = ""
    if data.tva_taux == 0:
        mention_tva = data.mention_legale_tva or "TVA non applicable, article 293 B du Code général des impôts"
    
    dessiner_pied_page(c, width, data, mention_tva)
    try:
        c.save()
        print(f"✅ PDF facture sauvegardé: {filepath}")
    except Exception as e:
        print(f"❌ Erreur lors de la sauvegarde du PDF: {e}")
        raise
    
    return filepath, numero_facture, total_ht, total_ttc


# ==================== GÉNÉRATION WORD ====================

def set_cell_shading(cell, color):
    """Applique une couleur de fond à une cellule Word"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generer_word_devis(data: DevisRequest, numero_devis_force: Optional[str] = None) -> str:
    """Génère un devis au format Word"""
    # PRIORITÉ 1: Utiliser le numéro forcé (paramètre explicite)
    # PRIORITÉ 2: Utiliser le numéro fourni dans data.numero_devis
    # PRIORITÉ 3: Générer un nouveau numéro (ne devrait jamais arriver)
    
    if numero_devis_force and str(numero_devis_force).strip():
        numero_devis = str(numero_devis_force).strip()
        print(f"✅ Word - Utilisation du numéro de devis FORCÉ (paramètre): '{numero_devis}'")
    elif data.numero_devis and str(data.numero_devis).strip():
        numero_devis = str(data.numero_devis).strip()
        print(f"✅ Word - Utilisation du numéro de devis fourni dans data: '{numero_devis}'")
    else:
        numero_devis = f"DEV-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"
        print(f"⚠️ Word - numero_devis non fourni ou vide, génération d'un nouveau numéro: {numero_devis}")
    
    filename = f"{numero_devis}.docx"
    filepath = os.path.join(PDF_FOLDER, filename)
    
    date_devis = datetime.now().strftime("%d/%m/%Y")
    date_validite = (datetime.now() + timedelta(days=data.validite_jours)).strftime("%d/%m/%Y")
    
    doc = Document()
    
    # Marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Logo si disponible
    logo_bytes = telecharger_logo_bytes(data.entreprise.logo_url)
    if logo_bytes:
        try:
            doc.add_picture(logo_bytes, width=Inches(1.2))
        except:
            pass
    
    # En-tête entreprise
    titre = doc.add_heading(data.entreprise.nom.upper(), 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in titre.runs:
        run.font.color.rgb = get_couleur_principale_rgb(data)
    
    if data.entreprise.gerant:
        p = doc.add_paragraph(f"Gérant : {data.entreprise.gerant}")
        p.runs[0].font.size = Pt(10)
    
    # DEVIS + Numéro
    doc.add_paragraph()
    titre_devis = doc.add_heading("DEVIS", 1)
    titre_devis.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p = doc.add_paragraph(f"N° {numero_devis}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(f"Date : {date_devis}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(f"Validité : {date_validite}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Tableau infos émetteur/destinataire
    table_info = doc.add_table(rows=1, cols=2)
    table_info.autofit = True
    
    # Émetteur
    cell_emetteur = table_info.rows[0].cells[0]
    cell_emetteur.text = ""
    p = cell_emetteur.add_paragraph()
    run = p.add_run("ÉMETTEUR")
    run.bold = True
    run.font.color.rgb = get_couleur_principale_rgb(data)
    cell_emetteur.add_paragraph(data.entreprise.nom)
    cell_emetteur.add_paragraph(data.entreprise.adresse)
    if data.entreprise.cp_ville:
        cell_emetteur.add_paragraph(data.entreprise.cp_ville)
    cell_emetteur.add_paragraph(f"Tél : {data.entreprise.tel}")
    cell_emetteur.add_paragraph(f"Email : {data.entreprise.email}")
    cell_emetteur.add_paragraph(f"SIRET : {data.entreprise.siret}")
    
    # Destinataire
    cell_dest = table_info.rows[0].cells[1]
    cell_dest.text = ""
    p = cell_dest.add_paragraph()
    run = p.add_run("DESTINATAIRE")
    run.bold = True
    run.font.color.rgb = get_couleur_principale_rgb(data)
    cell_dest.add_paragraph(data.client.nom)
    if data.client.adresse:
        cell_dest.add_paragraph(data.client.adresse)
    if data.client.cp_ville:
        cell_dest.add_paragraph(data.client.cp_ville)
    if data.client.tel:
        cell_dest.add_paragraph(f"Tél : {data.client.tel}")
    
    doc.add_paragraph()
    
    # Tableau des prestations
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # En-tête
    header_cells = table.rows[0].cells
    headers = ['Description', 'Qté', 'Unité', 'P.U. HT', 'Total HT']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], get_couleur_principale_hex_string(data))
    
    # Lignes
    total_ht = 0
    for prestation in data.prestations:
        row_cells = table.add_row().cells
        total_ligne = prestation.quantite * prestation.prix_unitaire
        total_ht += total_ligne
        
        row_cells[0].text = prestation.description
        row_cells[1].text = str(prestation.quantite)
        row_cells[2].text = prestation.unite
        row_cells[3].text = f"{prestation.prix_unitaire:.2f} €"
        row_cells[4].text = f"{total_ligne:.2f} €"
    
    doc.add_paragraph()
    
    # Totaux
    montant_tva = total_ht * (data.tva_taux / 100)
    total_ttc = total_ht + montant_tva
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Total HT : {total_ht:.2f} €")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if data.tva_taux > 0:
        p.add_run(f"TVA ({data.tva_taux}%) : {montant_tva:.2f} €")
    else:
        run = p.add_run("TVA non applicable")
        run.italic = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"TOTAL TTC : {total_ttc:.2f} €")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = get_couleur_principale_rgb(data)
    
    doc.add_paragraph()
    
    # Conditions
    doc.add_heading("CONDITIONS", 2)
    doc.add_paragraph(f"• Délai de réalisation : {data.delai_realisation}")
    doc.add_paragraph(f"• Conditions de paiement : {data.entreprise.conditions_paiement or data.conditions_paiement}")
    doc.add_paragraph(f"• Devis valable jusqu'au : {date_validite}")
    
    doc.add_paragraph()
    
    # Signature
    doc.add_paragraph("Bon pour accord")
    doc.add_paragraph("Date : ________________")
    doc.add_paragraph("Signature : ________________")
    
    # Pied de page
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{data.entreprise.nom} - SIRET {data.entreprise.siret}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    if data.tva_taux == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TVA non applicable, article 293 B du Code général des impôts")
        run.font.size = Pt(8)
        run.italic = True
    
    doc.save(filepath)
    
    return filepath, numero_devis, total_ht, total_ttc


def generer_word_facture(data: FactureRequest, numero_facture_force: Optional[str] = None) -> str:
    """Génère une facture au format Word"""
    # PRIORITÉ 1: Utiliser le numéro forcé (paramètre explicite)
    # PRIORITÉ 2: Utiliser le numéro fourni dans data.numero_facture
    # PRIORITÉ 3: Générer un nouveau numéro (ne devrait jamais arriver)
    
    if numero_facture_force and str(numero_facture_force).strip():
        numero_facture = str(numero_facture_force).strip()
        print(f"✅ Facture Word - Utilisation du numéro FORCÉ (paramètre): '{numero_facture}'")
    elif data.numero_facture and str(data.numero_facture).strip():
        numero_facture = str(data.numero_facture).strip()
        print(f"✅ Facture Word - Utilisation du numéro fourni dans data: '{numero_facture}'")
    else:
        numero_facture = f"FAC-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"
        print(f"⚠️ Facture Word - numero_facture non fourni ou vide, génération d'un nouveau numéro: {numero_facture}")
    
    filename = f"{numero_facture}.docx"
    filepath = os.path.join(PDF_FOLDER, filename)
    
    date_facture = datetime.now().strftime("%d/%m/%Y")
    date_echeance = (datetime.now() + timedelta(days=data.date_echeance_jours)).strftime("%d/%m/%Y")
    
    doc = Document()
    
    # Marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Logo si disponible
    logo_bytes = telecharger_logo_bytes(data.entreprise.logo_url)
    if logo_bytes:
        try:
            doc.add_picture(logo_bytes, width=Inches(1.2))
        except:
            pass
    
    # En-tête entreprise
    titre = doc.add_heading(data.entreprise.nom.upper(), 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in titre.runs:
        run.font.color.rgb = get_couleur_principale_rgb(data)
    
    if data.entreprise.gerant:
        p = doc.add_paragraph(f"Gérant : {data.entreprise.gerant}")
        p.runs[0].font.size = Pt(10)
    
    # FACTURE + Numéro
    doc.add_paragraph()
    titre_facture = doc.add_heading("FACTURE", 1)
    titre_facture.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in titre_facture.runs:
        run.font.color.rgb = get_couleur_principale_rgb(data)
    
    p = doc.add_paragraph(f"N° {numero_facture}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(f"Date : {date_facture}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if data.numero_devis_origine:
        p = doc.add_paragraph(f"Réf. devis : {data.numero_devis_origine}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(f"Échéance : {date_echeance}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Tableau infos émetteur/destinataire
    table_info = doc.add_table(rows=1, cols=2)
    table_info.autofit = True
    
    # Émetteur
    cell_emetteur = table_info.rows[0].cells[0]
    cell_emetteur.text = ""
    p = cell_emetteur.add_paragraph()
    run = p.add_run("ÉMETTEUR")
    run.bold = True
    run.font.color.rgb = get_couleur_principale_rgb(data)
    cell_emetteur.add_paragraph(data.entreprise.nom)
    cell_emetteur.add_paragraph(data.entreprise.adresse)
    if data.entreprise.cp_ville:
        cell_emetteur.add_paragraph(data.entreprise.cp_ville)
    cell_emetteur.add_paragraph(f"Tél : {data.entreprise.tel}")
    cell_emetteur.add_paragraph(f"Email : {data.entreprise.email}")
    cell_emetteur.add_paragraph(f"SIRET : {data.entreprise.siret}")
    
    # Destinataire
    cell_dest = table_info.rows[0].cells[1]
    cell_dest.text = ""
    p = cell_dest.add_paragraph()
    run = p.add_run("DESTINATAIRE")
    run.bold = True
    run.font.color.rgb = get_couleur_principale_rgb(data)
    cell_dest.add_paragraph(data.client.nom)
    if data.client.adresse:
        cell_dest.add_paragraph(data.client.adresse)
    if data.client.cp_ville:
        cell_dest.add_paragraph(data.client.cp_ville)
    if data.client.tel:
        cell_dest.add_paragraph(f"Tél : {data.client.tel}")
    if data.client.email:
        cell_dest.add_paragraph(f"Email : {data.client.email}")
    
    doc.add_paragraph()
    
    # Tableau des prestations
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # En-tête
    header_cells = table.rows[0].cells
    headers = ['Description', 'Qté', 'Unité', 'P.U. HT', 'Total HT']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], get_couleur_principale_hex_string(data))
    
    # Lignes
    total_ht = 0
    for prestation in data.prestations:
        row_cells = table.add_row().cells
        total_ligne = prestation.quantite * prestation.prix_unitaire
        total_ht += total_ligne
        
        row_cells[0].text = prestation.description
        row_cells[1].text = str(prestation.quantite)
        row_cells[2].text = prestation.unite
        row_cells[3].text = f"{prestation.prix_unitaire:.2f} €"
        row_cells[4].text = f"{total_ligne:.2f} €"
    
    doc.add_paragraph()
    
    # Totaux
    montant_tva = total_ht * (data.tva_taux / 100)
    total_ttc = total_ht + montant_tva
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Total HT : {total_ht:.2f} €")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if data.tva_taux > 0:
        p.add_run(f"TVA ({data.tva_taux}%) : {montant_tva:.2f} €")
    else:
        run = p.add_run("TVA non applicable")
        run.italic = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"TOTAL TTC : {total_ttc:.2f} €")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = get_couleur_principale_rgb(data)
    
    doc.add_paragraph()
    
    # Informations de paiement
    doc.add_heading("INFORMATIONS DE PAIEMENT", 2)
    doc.add_paragraph(f"• Date d'échéance : {date_echeance}")
    doc.add_paragraph("• Mode de paiement : Virement bancaire, chèque ou espèces")
    doc.add_paragraph("• En cas de retard : pénalité de 3 fois le taux d'intérêt légal")
    doc.add_paragraph("• Indemnité forfaitaire pour frais de recouvrement : 40€")
    
    # Pied de page
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{data.entreprise.nom} - SIRET {data.entreprise.siret}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    if data.tva_taux == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.mention_legale_tva or "TVA non applicable, article 293 B du Code général des impôts")
        run.font.size = Pt(8)
        run.italic = True
    
    doc.save(filepath)
    
    return filepath, numero_facture, total_ht, total_ttc


# ==================== ROUTES API ====================

@app.get("/")
def root():
    return {"message": "MonDevisPro API", "version": "3.0.0", "status": "ok"}


@app.post("/generer-devis")
async def generer_devis_endpoint(data: DevisRequest):
    try:
        # IMPORTANT: Récupérer le numéro AVANT toute autre opération
        # Si Pydantic n'a pas reçu le champ, il sera None
        numero_devis_recu = None
        
        # Essayer de récupérer depuis data.numero_devis
        if hasattr(data, 'numero_devis') and data.numero_devis:
            numero_devis_recu = str(data.numero_devis).strip()
            print(f"✅ Numéro de devis récupéré depuis data.numero_devis: '{numero_devis_recu}'")
        else:
            print(f"❌ ERREUR: data.numero_devis est None ou vide!")
            print(f"   - data.numero_devis = '{data.numero_devis}'")
            print(f"   - Type: {type(data.numero_devis)}")
            print(f"   - hasattr(data, 'numero_devis'): {hasattr(data, 'numero_devis')}")
            raise HTTPException(status_code=400, detail="Le numéro de devis est obligatoire et n'a pas été fourni dans la requête")
        
        if not numero_devis_recu or not numero_devis_recu.strip():
            print(f"❌ ERREUR CRITIQUE: Numéro de devis vide après traitement!")
            raise HTTPException(status_code=400, detail="Le numéro de devis est obligatoire")
        
        print(f"📄 Début génération devis pour client: {data.client.nom}")
        print(f"📊 Nombre de prestations: {len(data.prestations)}")
        print(f"🎨 Couleur PDF: {data.entreprise.couleur_pdf or 'défaut'}")
        print(f"🏢 Forme juridique: {data.entreprise.forme_juridique or 'non définie'}")
        print(f"💰 Capital social: {data.entreprise.capital_social or 'non défini'}")
        print(f"📋 RCS: {data.entreprise.rcs or 'non défini'}")
        print(f"📋 Numéro de devis à utiliser: '{numero_devis_recu}'")
        print(f"💰 Remise - type: '{data.remise_type}', valeur: {data.remise_valeur}, type valeur: {type(data.remise_valeur)}")
        
        # FORCER l'utilisation du numéro reçu en mettant à jour data.numero_devis
        # Utiliser model_copy pour Pydantic v2 ou copy pour v1
        try:
            if hasattr(data, 'model_copy'):
                data = data.model_copy(update={'numero_devis': numero_devis_recu})
            else:
                data.numero_devis = numero_devis_recu
            print(f"✅ data.numero_devis mis à jour avec: '{data.numero_devis}'")
        except Exception as e:
            print(f"⚠️ Impossible de mettre à jour data.numero_devis: {e}")
            # Créer un nouveau dict avec le numéro forcé
            data_dict = data.model_dump() if hasattr(data, 'model_dump') else data.dict()
            data_dict['numero_devis'] = numero_devis_recu
            data = DevisRequest(**data_dict)
            print(f"✅ data recréé avec numero_devis: '{data.numero_devis}'")
        
        # Générer PDF avec le numéro FORCÉ (paramètre explicite)
        print("📝 Génération PDF...")
        filepath_pdf, numero_devis_pdf, total_ht, total_ttc = generer_pdf_devis(data, numero_devis_force=numero_devis_recu)
        print(f"✅ PDF généré: {filepath_pdf}")
        print(f"📋 Numéro de devis utilisé dans PDF: '{numero_devis_pdf}'")
        print(f"📋 Numéro de devis reçu initialement: '{numero_devis_recu}'")
        
        # Le numéro utilisé DOIT correspondre au numéro reçu (on l'a forcé)
        if numero_devis_pdf != numero_devis_recu:
            print(f"❌ ERREUR CRITIQUE: Le numéro utilisé ({numero_devis_pdf}) diffère du numéro reçu ({numero_devis_recu})")
            print(f"   - Le PDF contient probablement le mauvais numéro!")
            # Renommer le fichier PDF pour correspondre au bon numéro
            correct_pdf_path = os.path.join(PDF_FOLDER, f"{numero_devis_recu}.pdf")
            if os.path.exists(filepath_pdf) and filepath_pdf != correct_pdf_path:
                print(f"🔄 Renommage du PDF de '{filepath_pdf}' vers '{correct_pdf_path}'")
                os.rename(filepath_pdf, correct_pdf_path)
                filepath_pdf = correct_pdf_path
            numero_devis_final = numero_devis_recu
        else:
            numero_devis_final = numero_devis_pdf
            print(f"✅ Le numéro de devis est cohérent: {numero_devis_final}")
        
        # Générer Word avec le numéro FORCÉ (paramètre explicite)
        print("📝 Génération Word...")
        filepath_word, numero_devis_word, _, _ = generer_word_devis(data, numero_devis_force=numero_devis_recu)
        # Renommer le Word pour avoir le même numéro que le PDF
        new_word_path = os.path.join(PDF_FOLDER, f"{numero_devis_final}.docx")
        if os.path.exists(filepath_word) and filepath_word != new_word_path:
            os.rename(filepath_word, new_word_path)
        print(f"✅ Word généré: {new_word_path}")
        
        # Upload sur Supabase Storage
        print("📤 Upload PDF sur Supabase...")
        pdf_url = upload_to_supabase(filepath_pdf, f"{numero_devis_final}.pdf")
        print(f"✅ PDF uploadé: {pdf_url}")
        
        print("📤 Upload Word sur Supabase...")
        word_url = upload_to_supabase(new_word_path, f"{numero_devis_final}.docx")
        print(f"✅ Word uploadé: {word_url}")
        
        return {
            "success": True,
            "numero_devis": numero_devis_final,  # IMPORTANT: Retourner le numéro final (celui du dashboard)
            "total_ht": total_ht,
            "total_ttc": total_ttc,
            "pdf_filename": f"{numero_devis_final}.pdf",
            "pdf_url": pdf_url,
            "word_filename": f"{numero_devis_final}.docx",
            "word_url": word_url
        }
    except Exception as e:
        print(f"❌ Erreur dans generer_devis_endpoint: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generer-devis-simple")
async def generer_devis_simple_endpoint(data: DevisRequestSimple):
    try:
        tva_taux = data.entreprise.tva_taux if data.entreprise.tva_taux is not None else 20.0
        conditions = data.entreprise.conditions_paiement or "30% à la commande, solde à réception"
        
        # Extraire les donnees client
        client_adresse = getattr(data.devis_data, 'client_adresse', '') or ''
        client_email = getattr(data.devis_data, 'client_email', '') or ''
        client_telephone = getattr(data.devis_data, 'client_telephone', '') or ''
        acompte = getattr(data.devis_data, 'acompte_pourcentage', 0) or 0
        
        # Gerer les prestations: soit liste directe, soit JSON string
        prestations_list = data.devis_data.prestations
        if not prestations_list and data.devis_data.prestations_json:
            try:
                import json
                from urllib.parse import unquote
                # Decoder l'URL encoding si present
                json_str = unquote(data.devis_data.prestations_json)
                print(f"📋 Prestations JSON decodee: {json_str[:200]}...")
                parsed = json.loads(json_str)
                prestations_list = [Prestation(**p) for p in parsed]
                print(f"✅ Prestations parsees depuis JSON string: {len(prestations_list)} lignes")
            except Exception as e:
                print(f"❌ Erreur parsing prestations_json: {e}")
                prestations_list = []
        
        if not prestations_list:
            return {"success": False, "error": "Aucune prestation fournie"}
        
        full_data = DevisRequest(
            entreprise=data.entreprise,
            client=Client(
                nom=data.devis_data.client_nom,
                adresse=client_adresse,
                cp_ville="",
                tel=client_telephone,
                email=client_email
            ),
            prestations=prestations_list,
            tva_taux=tva_taux,
            conditions_paiement=conditions,
            delai_realisation=data.devis_data.delai,
            validite_jours=data.validite_jours,
            remise_type=data.devis_data.remise_type,
            remise_valeur=data.devis_data.remise_valeur or 0,
            acompte_pourcentage=acompte,
            numero_devis=None  # Pour l'IA, on peut generer un nouveau numero
        )
        
        # Générer PDF
        filepath_pdf, numero_devis, total_ht, total_ttc = generer_pdf_devis(full_data)
        
        # Générer Word
        filepath_word, _, _, _ = generer_word_devis(full_data)
        new_word_path = os.path.join(PDF_FOLDER, f"{numero_devis}.docx")
        if os.path.exists(filepath_word) and filepath_word != new_word_path:
            os.rename(filepath_word, new_word_path)
        
        # Upload sur Supabase Storage
        pdf_url = upload_to_supabase(filepath_pdf, f"{numero_devis}.pdf")
        word_url = upload_to_supabase(new_word_path, f"{numero_devis}.docx")
        
        # ============================================================
        # SAUVEGARDE AUTOMATIQUE DANS LE DASHBOARD (si phone fourni)
        # ============================================================
        devis_dashboard_id = None
        if data.phone:
            print(f"📱 Phone fourni: {data.phone} - Recherche entreprise...")
            entreprise = get_entreprise_by_whatsapp(data.phone)
            if entreprise:
                # Préparer les prestations pour le dashboard
                prestations_for_db = []
                for p in prestations_list:
                    prestations_for_db.append({
                        'description': p.description,
                        'quantite': p.quantite,
                        'unite': p.unite,
                        'prix_unitaire_ht': p.prix_unitaire,
                        'prix_unitaire': p.prix_unitaire,
                        'tva_taux': p.tva_taux if p.tva_taux else tva_taux,
                    })
                
                # Sauvegarder dans le dashboard
                saved_devis = save_devis_to_dashboard(
                    entreprise_id=entreprise['id'],
                    numero_devis=numero_devis,
                    client_nom=data.devis_data.client_nom,
                    client_email=client_email,
                    client_telephone=client_telephone,
                    titre_projet=getattr(data.devis_data, 'titre_projet', None),
                    prestations=prestations_for_db,
                    total_ht=total_ht,
                    total_ttc=total_ttc,
                    pdf_url=pdf_url,
                    word_url=word_url,
                    remise_type=data.devis_data.remise_type,
                    remise_value=data.devis_data.remise_valeur,
                    delai=data.devis_data.delai
                )
                if saved_devis:
                    devis_dashboard_id = saved_devis.get('id')
                    print(f"✅ Devis sauvegardé dans dashboard avec ID: {devis_dashboard_id}")
            else:
                print(f"⚠️ Entreprise non trouvée pour {data.phone} - Devis non sauvegardé dans dashboard")
        
        return {
            "success": True,
            "numero_devis": numero_devis,
            "total_ht": total_ht,
            "total_ttc": total_ttc,
            "pdf_filename": f"{numero_devis}.pdf",
            "pdf_url": pdf_url,
            "word_filename": f"{numero_devis}.docx",
            "word_url": word_url,
            "dashboard_id": devis_dashboard_id  # ID dans le dashboard (si sauvegardé)
        }
    except Exception as e:
        print(f"❌ Erreur dans generer_devis_simple_endpoint: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generer-facture")
async def generer_facture_endpoint(data: FactureRequest):
    try:
        # IMPORTANT: Récupérer le numéro AVANT toute autre opération
        # Si Pydantic n'a pas reçu le champ, il sera None
        numero_facture_recu = None
        
        # Essayer de récupérer depuis data.numero_facture
        if hasattr(data, 'numero_facture') and data.numero_facture:
            numero_facture_recu = str(data.numero_facture).strip()
            print(f"✅ Numéro de facture récupéré depuis data.numero_facture: '{numero_facture_recu}'")
        else:
            # Si le numéro n'est pas fourni, générer un numéro par défaut (pour rétrocompatibilité)
            # mais logger un avertissement
            numero_facture_recu = f"FAC-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"
            print(f"⚠️ AVERTISSEMENT: data.numero_facture est None ou vide!")
            print(f"   - data.numero_facture = '{data.numero_facture}'")
            print(f"   - Type: {type(data.numero_facture)}")
            print(f"   - hasattr(data, 'numero_facture'): {hasattr(data, 'numero_facture')}")
            print(f"   - Génération d'un numéro par défaut: '{numero_facture_recu}'")
            print(f"   - ⚠️ Ce numéro pourrait ne pas correspondre au numéro en base de données!")
        
        if not numero_facture_recu or not numero_facture_recu.strip():
            # Dernière vérification de sécurité
            numero_facture_recu = f"FAC-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"
            print(f"⚠️ Numéro de facture vide après traitement, génération d'un numéro par défaut: '{numero_facture_recu}'")
        
        # Parser prestations_json si fourni (pour Make.com)
        if getattr(data, 'prestations_json', None) and (not data.prestations or len(data.prestations) == 0):
            try:
                import json
                from urllib.parse import unquote
                json_str = unquote(data.prestations_json)
                print(f"📋 Prestations JSON decodee: {json_str[:200]}...")
                parsed = json.loads(json_str)
                prestations_list = [Prestation(**p) for p in parsed]
                # Mettre à jour data avec les nouvelles prestations
                if hasattr(data, 'model_copy'):
                    data = data.model_copy(update={'prestations': prestations_list})
                else:
                    data.prestations = prestations_list
                print(f"✅ Prestations parsees depuis JSON string: {len(prestations_list)} lignes")
            except Exception as e:
                print(f"❌ Erreur parsing prestations_json: {e}")
        
        print(f"📄 Début génération facture pour client: {data.client.nom}")
        print(f"📊 Nombre de prestations: {len(data.prestations)}")
        print(f"🎨 Couleur PDF: {data.entreprise.couleur_pdf or 'défaut'}")
        print(f"🏢 Forme juridique: {data.entreprise.forme_juridique or 'non définie'}")
        print(f"💰 Capital social: {data.entreprise.capital_social or 'non défini'}")
        print(f"📋 RCS: {data.entreprise.rcs or 'non défini'}")
        print(f"📋 Numéro de facture à utiliser: '{numero_facture_recu}'")
        
        # DEBUG: Vérifier les valeurs pour facture d'acompte
        is_facture_acompte = getattr(data, 'is_facture_acompte', False)
        taux_acompte = getattr(data, 'taux_acompte', None)
        total_ttc_recu = getattr(data, 'total_ttc', None)
        total_ht_recu = getattr(data, 'total_ht', None)
        print(f"🔍 DEBUG FACTURE ACOMPTE:")
        print(f"   is_facture_acompte: {is_facture_acompte}")
        print(f"   taux_acompte: {taux_acompte}")
        print(f"   total_ttc reçu: {total_ttc_recu} (type: {type(total_ttc_recu)})")
        print(f"   total_ht reçu: {total_ht_recu} (type: {type(total_ht_recu)})")
        if data.prestations and len(data.prestations) > 0:
            print(f"   prix_unitaire prestation: {data.prestations[0].prix_unitaire}")
            print(f"   quantite prestation: {data.prestations[0].quantite}")
        
        # ============================================================
        # CALCUL AUTOMATIQUE DU MONTANT D'ACOMPTE SI taux_acompte fourni
        # ============================================================
        if is_facture_acompte and taux_acompte and taux_acompte > 0:
            # PRIORITÉ : Utiliser total_ht_devis/total_ttc_devis (inclut la remise)
            total_ht_devis = getattr(data, 'total_ht_devis', None)
            total_ttc_devis = getattr(data, 'total_ttc_devis', None)
            
            if total_ht_devis and total_ttc_devis:
                # Utiliser les totaux du devis (avec remise déjà appliquée)
                print(f"📊 UTILISATION DES TOTAUX DU DEVIS (avec remise):")
                print(f"   Total HT devis: {total_ht_devis}")
                print(f"   Total TTC devis: {total_ttc_devis}")
                total_ht_base = total_ht_devis
                total_ttc_base = total_ttc_devis
            else:
                # Fallback : Calculer à partir des prestations (sans remise)
                tva_taux = getattr(data.entreprise, 'tva_taux', 20) or 20
                total_ht_base = 0
                for p in data.prestations:
                    total_ht_base += p.prix_unitaire * p.quantite
                total_ttc_base = total_ht_base * (1 + tva_taux / 100)
                print(f"⚠️ CALCUL DEPUIS PRESTATIONS (sans remise):")
                print(f"   Total HT calculé: {total_ht_base}")
                print(f"   Total TTC calculé: {total_ttc_base}")
            
            # Appliquer le taux d'acompte
            total_ht_acompte = round(total_ht_base * taux_acompte / 100, 2)
            total_ttc_acompte = round(total_ttc_base * taux_acompte / 100, 2)
            
            print(f"📊 CALCUL ACOMPTE:")
            print(f"   Taux acompte: {taux_acompte}%")
            print(f"   Total HT acompte: {total_ht_acompte}")
            print(f"   Total TTC acompte: {total_ttc_acompte}")
            
            # Mettre à jour les totaux dans data
            total_ttc_recu = total_ttc_acompte
            total_ht_recu = total_ht_acompte
            
            # Mettre à jour l'objet data avec les montants d'acompte
            try:
                if hasattr(data, 'model_copy'):
                    data = data.model_copy(update={'total_ttc': total_ttc_acompte, 'total_ht': total_ht_acompte})
                else:
                    data.total_ttc = total_ttc_acompte
                    data.total_ht = total_ht_acompte
                print(f"✅ Montants d'acompte appliqués: HT={total_ht_acompte}, TTC={total_ttc_acompte}")
            except Exception as e:
                print(f"⚠️ Erreur mise à jour montants acompte: {e}")
        
        # ============================================================
        # DÉTECTION AUTOMATIQUE DES FACTURES D'ACOMPTE
        # ============================================================
        # Si is_facture_acompte n'est pas explicitement True, on le détecte automatiquement
        if not is_facture_acompte:
            # Vérifier si le numéro de facture contient "ACO"
            if numero_facture_recu and "ACO" in numero_facture_recu.upper():
                is_facture_acompte = True
                print(f"✅ DÉTECTION AUTO: Facture d'acompte détectée via numéro '{numero_facture_recu}'")
            # Vérifier si la description contient "Acompte"
            elif data.prestations and len(data.prestations) == 1:
                desc = getattr(data.prestations[0], 'description', '')
                if 'acompte' in desc.lower():
                    is_facture_acompte = True
                    print(f"✅ DÉTECTION AUTO: Facture d'acompte détectée via description '{desc}'")
            # Vérifier si total_ttc est fourni et différent du calcul
            if total_ttc_recu is not None and total_ht_recu is not None:
                is_facture_acompte = True
                print(f"✅ DÉTECTION AUTO: Facture d'acompte détectée via total_ttc/total_ht fournis")
        
        # FORCER les mises à jour sur l'objet data (numéro + is_facture_acompte)
        updates = {'numero_facture': numero_facture_recu}
        if is_facture_acompte:
            updates['is_facture_acompte'] = True
        
        try:
            if hasattr(data, 'model_copy'):
                data = data.model_copy(update=updates)
            else:
                data.numero_facture = numero_facture_recu
                if is_facture_acompte:
                    data.is_facture_acompte = True
            print(f"✅ data mis à jour - numero_facture: '{data.numero_facture}', is_facture_acompte: {data.is_facture_acompte}")
            print(f"   total_ttc dans data: {data.total_ttc}, total_ht dans data: {data.total_ht}")
        except Exception as e:
            print(f"⚠️ Impossible de mettre à jour data: {e}")
            # Créer un nouveau dict avec les valeurs forcées
            data_dict = data.model_dump() if hasattr(data, 'model_dump') else data.dict()
            data_dict.update(updates)
            data = FactureRequest(**data_dict)
            print(f"✅ data recréé avec numero_facture: '{data.numero_facture}', is_facture_acompte: {data.is_facture_acompte}")
        
        # Générer PDF avec le numéro forcé
        filepath_pdf, numero_facture_pdf, total_ht, total_ttc = generer_pdf_facture(data, numero_facture_force=numero_facture_recu)
        
        # Vérifier que le numéro utilisé correspond bien au numéro reçu
        if numero_facture_pdf != numero_facture_recu:
            print(f"❌ ERREUR CRITIQUE: Le numéro utilisé ({numero_facture_pdf}) diffère du numéro reçu ({numero_facture_recu})")
            # Utiliser le numéro reçu (celui du dashboard) - c'est la source de vérité
            numero_facture_final = numero_facture_recu
            # Renommer le fichier PDF pour correspondre au bon numéro
            correct_pdf_path = os.path.join(PDF_FOLDER, f"{numero_facture_final}.pdf")
            if os.path.exists(filepath_pdf) and filepath_pdf != correct_pdf_path:
                print(f"🔄 Renommage du PDF de '{filepath_pdf}' vers '{correct_pdf_path}'")
                os.rename(filepath_pdf, correct_pdf_path)
                filepath_pdf = correct_pdf_path
        else:
            numero_facture_final = numero_facture_pdf
            print(f"✅ Le numéro de facture est cohérent: {numero_facture_final}")
        
        # Générer Word avec le numéro forcé
        filepath_word, _, _, _ = generer_word_facture(data, numero_facture_force=numero_facture_recu)
        new_word_path = os.path.join(PDF_FOLDER, f"{numero_facture_final}.docx")
        if os.path.exists(filepath_word) and filepath_word != new_word_path:
            print(f"🔄 Renommage du Word de '{filepath_word}' vers '{new_word_path}'")
            os.rename(filepath_word, new_word_path)
        
        # Upload sur Supabase Storage
        pdf_url = upload_to_supabase(filepath_pdf, f"{numero_facture_final}.pdf")
        word_url = upload_to_supabase(new_word_path, f"{numero_facture_final}.docx")
        
        # ============================================================
        # SAUVEGARDE AUTOMATIQUE DANS LE DASHBOARD (si phone fourni)
        # ============================================================
        facture_dashboard_id = None
        devis_id_for_facture = None
        
        if getattr(data, 'phone', None):
            print(f"📱 Phone fourni: {data.phone} - Recherche entreprise...")
            entreprise = get_entreprise_by_whatsapp(data.phone)
            if entreprise:
                # Si numero_devis_origine fourni, trouver le devis dans le dashboard
                if data.numero_devis_origine:
                    devis_existant = get_devis_by_numero(data.numero_devis_origine, entreprise['id'])
                    if devis_existant:
                        devis_id_for_facture = devis_existant.get('id')
                        print(f"✅ Devis trouvé: {devis_id_for_facture}")
                
                # Préparer les prestations pour le dashboard
                prestations_for_db = []
                if data.prestations:
                    for p in data.prestations:
                        prestations_for_db.append({
                            'description': p.description,
                            'quantite': p.quantite,
                            'unite': p.unite,
                            'prix_unitaire_ht': p.prix_unitaire,
                            'prix_unitaire': p.prix_unitaire,
                            'tva_taux': p.tva_taux if p.tva_taux else data.tva_taux,
                        })
                
                # Déterminer le type de facture
                type_facture = 'acompte' if is_facture_acompte else 'complete'
                
                # Sauvegarder dans le dashboard
                saved_facture = save_facture_to_dashboard(
                    entreprise_id=entreprise['id'],
                    devis_id=devis_id_for_facture,
                    numero_facture=numero_facture_final,
                    client_nom=data.client.nom,
                    client_email=data.client.email,
                    client_telephone=data.client.tel,
                    client_adresse=data.client.adresse,
                    titre_projet=None,  # On pourrait le récupérer du devis
                    prestations=prestations_for_db,
                    total_ht=total_ht,
                    total_ttc=total_ttc,
                    pdf_url=pdf_url,
                    word_url=word_url,
                    type_facture=type_facture,
                    remise_type=data.remise_type,
                    remise_value=data.remise_valeur,
                    tva_taux=data.tva_taux
                )
                if saved_facture:
                    facture_dashboard_id = saved_facture.get('id')
                    print(f"✅ Facture sauvegardée dans dashboard avec ID: {facture_dashboard_id}")
            else:
                print(f"⚠️ Entreprise non trouvée pour {data.phone} - Facture non sauvegardée dans dashboard")
        
        return {
            "success": True,
            "numero_facture": numero_facture_final,
            "total_ht": total_ht,
            "total_ttc": total_ttc,
            "pdf_filename": f"{numero_facture_final}.pdf",
            "pdf_url": pdf_url,
            "word_filename": f"{numero_facture_final}.docx",
            "word_url": word_url,
            "dashboard_id": facture_dashboard_id,  # ID dans le dashboard (si sauvegardé)
            "devis_id": devis_id_for_facture  # ID du devis lié (si trouvé)
        }
    except Exception as e:
        print(f"❌ Erreur dans generer_facture_endpoint: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/download/{filename}")
async def download_file(filename: str):
    filepath = os.path.join(PDF_FOLDER, filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="Fichier non trouvé")
    
    # Déterminer le type MIME
    if filename.endswith('.pdf'):
        media_type = "application/pdf"
    elif filename.endswith('.docx'):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        media_type = "application/octet-stream"
    
    return FileResponse(filepath, media_type=media_type, filename=filename)


@app.get("/health")
def health_check():
    return {"status": "healthy"}




# =============================================================================
# WHATSAPP HANDLER (module externe — whatsapp_handler.py)
# =============================================================================

from whatsapp_handler import router as whatsapp_router, setup as whatsapp_setup

whatsapp_setup({
    "supabase_client": supabase_client,
    "anthropic_client": anthropic_client,
    "openai_whisper_client": openai_whisper_client,
    "get_entreprise_by_whatsapp": get_entreprise_by_whatsapp,
    "save_devis_to_dashboard": save_devis_to_dashboard,
    "save_facture_to_dashboard": save_facture_to_dashboard,
    "generer_pdf_devis": generer_pdf_devis,
    "generer_word_devis": generer_word_devis,
    "generer_pdf_facture": generer_pdf_facture,
    "generer_word_facture": generer_word_facture,
    "upload_to_supabase": upload_to_supabase,
    "Prestation": Prestation,
    "Entreprise": Entreprise,
    "Client": Client,
    "DevisRequest": DevisRequest,
    "FactureRequest": FactureRequest,
})

app.include_router(whatsapp_router)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
